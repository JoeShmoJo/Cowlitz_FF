#make_memo_draft2_section8.py
# -*- coding: utf-8 -*-
"""
Write Section 8 and Appendix F into a COPY of the hand-edited combined memo.

WHY THIS IS NOT A COPY OF THE GENERATOR
    The first attempt at this copied make_memo_combined.py and re-ran it. That
    was wrong. The generator rebuilds the document from source every time, so
    regenerating threw away every hand edit made to the .docx since it was
    last generated -- 34 paragraphs of them, including a rewritten Purpose and
    substantial edits through Sections 2 and 4.

    This script never regenerates. It COPIES the edited .docx byte for byte
    and then edits only two places inside the copy:

        Section 8   the "[TO BE COMPLETED]" placeholder is replaced
        Appendix F  inserted between Appendix E and References

    Everything else in the document -- every hand edit, every table, every
    figure already placed -- is carried through untouched, because it is
    never rewritten.

    The corollary: hand edits made to DRAFT2 will be lost if this is re-run,
    because DRAFT2 is recreated from DRAFT each time. Edit DRAFT (the source
    of truth for prose) or stop running this once DRAFT2 is under review.

STYLING
    The paragraph, table, figure and caption helpers are imported from
    make_memo_combined rather than reimplemented, so anything inserted here
    formats identically to the rest of the document. Importing that module
    does not regenerate anything -- its main() is behind __name__ == guard.

INPUTS
    MEMO_CAS_Combined_FlowFrequency_DRAFT.docx   the hand-edited draft
    ../output/below_confluence_frequency.csv     Section 8 numbers
    ../output/diagnostics/prism_basin_precip_ratio.csv
    figures_combined/*.png

OUTPUT
    MEMO_CAS_Combined_FlowFrequency_DRAFT2.docx
"""

import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

import shutil

import pandas as pd
from docx import Document

# Styling helpers only. This import does NOT regenerate the memo.
import make_memo_combined as mm

# ----------------------------------------------------------------------------
# USER SETTINGS
# ----------------------------------------------------------------------------
SRC_DOCX = r"MEMO_CAS_Combined_FlowFrequency_DRAFT.docx"
OUT_DOCX = r"MEMO_CAS_Combined_FlowFrequency_DRAFT2.docx"

BELOW_CONF_CSV = r"../output/below_confluence_frequency.csv"
PRISM_CSV = r"../output/diagnostics/prism_basin_precip_ratio.csv"

# Headings used to find where to cut and where to insert. Matched on the
# leading text, so minor hand edits to the wording still resolve.
SECTION8_PREFIX = "8."
APPENDIX_A_PREFIX = "Appendix A"
REFERENCES_PREFIX = "References"

NEW_FIGURES = {
    "below_conf": r"figures_combined/below_confluence_frequency.png",
    "freq_gage": r"figures_combined/freq_castle_rock_gage.png",
    "freq_arkansas": r"figures_combined/freq_below_arkansas_creek.png",
    "freq_ostrander": r"figures_combined/freq_below_ostrander_creek.png",
    "freq_coweeman": r"figures_combined/freq_below_coweeman_river.png",
    "convergence": r"figures_combined/transform_convergence.png",
}

# ----------------------------------------------------------------------------


def find_heading(doc, prefix):
    """First Heading paragraph whose text starts with prefix."""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix):
            return p
    raise SystemExit("Could not find a heading starting with %r. The source "
                     "document's structure has changed; check it before "
                     "trusting this script." % prefix)


def delete_through(body, start_el, stop_el):
    """Remove start_el and everything after it, up to but not including stop_el."""
    n, el = 0, start_el
    while el is not None and el is not stop_el:
        nxt = el.getnext()
        body.remove(el)
        el = nxt
        n += 1
    return n


def capture(doc, build):
    """Run build(doc), which appends, and return the elements it appended.

    python-docx only adds at the end of the body, so content destined for the
    middle of the document is built at the end and then moved.
    """
    body = doc.element.body
    before = list(body)
    build(doc)
    return [el for el in body if el not in before]


def move_before(anchor_p, elements):
    for el in elements:
        anchor_p._p.addprevious(el)


# ----------------------------------------------------------------------------
# the inserted content
# ----------------------------------------------------------------------------


def build_section8(doc, conf, prism, tgt):
    mm.h1(doc, "8. Extension Downstream to the Coweeman Confluence")
    mm.para(doc, "The regulated frequency curve developed above applies at "
                 "the Castle Rock gage, which drains 2,229 square miles. "
                 "Between that gage and the Coweeman River confluence the "
                 "Cowlitz picks up 247 square miles of additional drainage — "
                 "Arkansas Creek, Ostrander Creek, the Coweeman River, and "
                 "the ungaged local area along both banks. This section "
                 "develops regulated peak frequency curves at each of those "
                 "points. No analysis below the Coweeman confluence is "
                 "included.")

    mm.h2(doc, "8.1 Method")
    mm.para(doc, "At a site draining DA square miles, the regulated peak of "
                 "annual exceedance probability p is taken as the regulated "
                 "Castle Rock peak of the same probability plus a local "
                 "contribution:")
    mm.para(doc, "Q(p) = Q_reg(p) + Q_unreg(p) × (DA − 2,229)/2,229 × 0.80",
            italic=True)
    mm.para(doc, "The local term scales off the UNREGULATED curve, not the "
                 "regulated one. All three tributaries enter below Mossyrock "
                 "and Riffe and respond to the storm rather than to the "
                 "release, so the storm-scale curve is what governs them; "
                 "scaling them off the regulated curve would shrink them in "
                 "proportion to a reservoir upstream of them, about 18 "
                 "percent low at the 1,000-year event.")
    mm.para(doc, "Pairing the unregulated and regulated curves at a common "
                 "probability is not a coincident-frequency assumption. It is "
                 "the same river in the same event, and the regulated curve "
                 "was derived from the unregulated one by routing (Section "
                 "5.4). Nothing in this method requires a tributary to be "
                 "simultaneously at its own 1,000-year event, which is the "
                 "assumption the earlier same-probability summation methods "
                 "carried and could not defend.")

    mm.h2(doc, "8.2 Local Contribution by Drainage Area")
    mm.para(doc, "Drainage areas here are taken from the StreamStats "
                 "delineations prepared for this analysis, which give 2,229 "
                 "square miles at the Castle Rock gage. Section 2 cites 2,238 "
                 "square miles from the earlier memoranda. The two differ by "
                 "0.4 percent and the difference is immaterial to a ratio, "
                 "but the delineated value is used throughout this section so "
                 "that every area in it comes from one consistent source.")
    mm.para(doc, "Each site uses the full incremental area between it and the "
                 "gage, not the named tributary basin alone. The difference "
                 "is ungaged local drainage along both banks, which "
                 "contributes whether or not it carries a name — 247 square "
                 "miles by incremental area against 197.5 for the three named "
                 "basins.")
    mm.table(doc,
             ["Location", "Drainage area (sq mi)", "Incremental (sq mi)",
              "Local as % of gage"],
             [["Castle Rock gage", "2,229", "—", "—"],
              ["Below Arkansas Creek", "2,278", "49", "2.2%"],
              ["Below Ostrander Creek", "2,335", "106", "4.8%"],
              ["Below Coweeman River", "2,476", "247", "11.1%"]],
             align_right=(1, 2, 3))
    mm.caption(doc, "Table 8-1. Drainage area at each location.")
    mm.para(doc, "A plain drainage-area ratio is adopted, with no unit-runoff "
                 "adjustment. Three independent lines of evidence support "
                 "that at the magnitudes that matter. Paired against the "
                 "routed unregulated Castle Rock peak for the 23 water years "
                 "in which both rivers peaked in the same storm, the Coweeman "
                 "gage record (USGS 14245000, WY1950–1996) runs about 1.5 "
                 "times its area share at common events but converges "
                 "monotonically toward parity as events grow; in February "
                 "1996 — a 212,245 cfs unregulated event, 92 percent of the "
                 "1,000-year flow — it sat at 1.04 times its area share. The "
                 "shorter Ecology record at gage 26C075 (WY2007–2019) gives "
                 "1.11 times for its largest bin, and that figure is a lower "
                 "bound because the gage's rating ceiling censors precisely "
                 "its largest events.")
    mm.para(doc, "Basin mean precipitation tests the equal-depth assumption "
                 "the area ratio rests on directly. Over the %d years %d–%d, "
                 "PRISM annual precipitation averaged across the delineated "
                 "boundaries gives a Coweeman-to-Castle-Rock ratio of %.2f "
                 "(median %.2f, interquartile range %.2f–%.2f). The two "
                 "basins receive very nearly the same depth."
                 % (len(prism), int(prism["year"].min()),
                    int(prism["year"].max()), prism["ratio"].mean(),
                    prism["ratio"].median(), prism["ratio"].quantile(0.25),
                    prism["ratio"].quantile(0.75)))
    mm.para(doc, "Read together these lines resolve into one consistent "
                 "picture. Precipitation is essentially equal per unit area, "
                 "so the Coweeman's excess at common events is a response "
                 "difference — a small, low, steep basin concentrating runoff "
                 "faster than a 2,229 square mile mainstem — rather than a "
                 "precipitation difference. Response differences shrink as "
                 "both basins saturate, which is why the flow ratio falls "
                 "toward parity as events grow, and why at the largest event "
                 "on record it converges on very nearly the precipitation "
                 "ratio itself. A plain area ratio is therefore the right "
                 "form at the magnitudes this study is concerned with, even "
                 "though it would understate the tributary at ordinary flows.")

    mm.h2(doc, "8.3 Timing Adjustment")
    mm.para(doc, "The tributaries do not crest when the regulated Cowlitz "
                 "does. They peak first and are already receding when the "
                 "regulated mainstem arrives, so only part of each tributary "
                 "peak is present at the moment that governs the combined "
                 "flow. Measuring the Coweeman flow at the hour of the "
                 "regulated Castle Rock crest, as a fraction of its own peak "
                 "in the same storm, over 78 events:")
    mm.table(doc,
             ["Castle Rock unregulated peak", "Events", "Median ratio"],
             [["20,000–40,000 cfs", "52", "0.806"],
              ["40,000–60,000 cfs", "19", "0.781"],
              ["Above 60,000 cfs", "7", "0.413"],
              ["All events", "78", "0.789"]],
             align_right=(1, 2))
    mm.caption(doc, "Table 8-2. Coweeman flow at the regulated Castle Rock "
                    "crest, as a fraction of its own event peak.")
    mm.para(doc, "A factor of 0.80 is adopted. It is where the two "
                 "well-sampled bins group, and it is the value used in the "
                 "2009 restudy, so it carries precedent as well as data. It "
                 "is very likely conservative: a higher timing factor adds "
                 "more local flow, and the seven largest events in the record "
                 "sit well below it.")
    mm.para(doc, "The lower figure from those seven largest events is not "
                 "adopted, for three reasons. The sample is small. Its "
                 "mechanism is tributary lead time rather than event "
                 "magnitude — the ratio correlates strongly with how far "
                 "ahead the Coweeman crests (Spearman rho −0.55, p < 0.0001) "
                 "while lead time itself shows no significant relationship to "
                 "event size (rho −0.04, p = 0.73), so the low tail figure is "
                 "a property of those particular storms rather than of large "
                 "events generally. And three events whose Coweeman crest "
                 "exceeded the gage rating are absent from that bin "
                 "altogether — events in which the tributary was at its "
                 "largest, so their absence biases the figure downward.")

    mm.h2(doc, "8.4 Results")
    mm.para(doc, "Regulated peak flows at the four locations are given below "
                 "for the probabilities of primary interest; complete "
                 "ordinates with confidence limits are in Appendix F.")
    rows = []
    for aep in [0.10, 0.02, 0.01, 0.005, 0.002, 0.001]:
        r = conf.iloc[(conf["AEP"] - aep).abs().idxmin()]
        rows.append(["%g%%" % (r["AEP"] * 100),
                     mm.fmt(r["castle_rock_gage_cfs"]),
                     mm.fmt(r["below_arkansas_creek_cfs"]),
                     mm.fmt(r["below_ostrander_creek_cfs"]),
                     mm.fmt(r["below_coweeman_river_cfs"])])
    mm.table(doc, ["AEP", "Castle Rock gage", "Below Arkansas",
                   "Below Ostrander", "Below Coweeman"], rows,
             align_right=(1, 2, 3, 4))
    mm.caption(doc, "Table 8-3. Regulated peak flow at each location, cfs.")
    mm.para(doc, "At the 1,000-year event the local contribution below the "
                 "Coweeman confluence is %s cfs, raising the regulated peak "
                 "%.1f percent above the gage. The increase is "
                 "proportionally largest in the middle of the range, near "
                 "the 2 percent event, and falls away in the extreme tail: "
                 "the regulated curve steepens there as the projects lose "
                 "the ability to hold back a rare inflow, while the local "
                 "term grows only with the flatter unregulated curve."
                 % (mm.fmt(tgt["below_coweeman_river_local_cfs"]),
                    100 * (tgt["below_coweeman_river_cfs"]
                           - tgt["castle_rock_gage_cfs"])
                    / tgt["castle_rock_gage_cfs"]))
    mm.figure(doc, "below_conf",
              "Figure 8-1. Regulated peak flow frequency at all four "
              "locations. The local contribution at each site is the "
              "incremental drainage area scaled off the unregulated curve and "
              "reduced by the 0.80 timing factor. The lower panel gives each "
              "site as a percentage increase over the gage.")
    mm.figure(doc, "freq_gage",
              "Figure 8-2. Castle Rock gage, 2,229 square miles: the "
              "unregulated and regulated curves of Sections 4 and 5, with the "
              "regulated 95 percent confidence band from Section 5.6. This is "
              "the pair every downstream location is built from — the "
              "regulated curve sets the base and the unregulated curve drives "
              "the local contribution.")
    mm.figure(doc, "freq_arkansas",
              "Figure 8-3. Below Arkansas Creek, 2,278 square miles.")
    mm.figure(doc, "freq_ostrander",
              "Figure 8-4. Below Ostrander Creek, 2,335 square miles.")
    mm.figure(doc, "freq_coweeman",
              "Figure 8-5. Below the Coweeman River, 2,476 square miles.")
    mm.para(doc, "The confidence band shown at each downstream location is "
                 "the Castle Rock regulated band of Section 5.6, carried "
                 "forward and translated by that site's local contribution. "
                 "No additional uncertainty is added for the local term "
                 "itself. That term carries real uncertainty, but it is a few "
                 "percent of a quantity whose own band is already wider than "
                 "the flow it brackets, so adding it would not be visible on "
                 "these figures and would imply a precision this method does "
                 "not have. The band below the Coweeman confluence spans %s "
                 "to %s cfs at the 1,000-year event. The choice of timing "
                 "factor is small against that: the full 0.41 to 1.00 range "
                 "moves the total 7.2 percent, about 6 percent of the width "
                 "of the confidence band."
                 % (mm.fmt(tgt["below_coweeman_river_lower_cfs"]),
                    mm.fmt(tgt["below_coweeman_river_upper_cfs"])))

    mm.h2(doc, "8.5 Behavior of the Transform at Extreme Flows")
    mm.para(doc, "A reservoir cannot remove volume from a flood, only move it "
                 "in time — at the 5-day duration the regulated to "
                 "unregulated ratio already exceeds one (Section 5.2). Riffe "
                 "Lake holds 358,116 acre-feet between the winter rule curve "
                 "at 745.5 feet and full pool at 778.5 feet, which is 180,550 "
                 "cfs-days. Once an event's pre-crest inflow fills that, "
                 "every further cubic foot per second passes through and the "
                 "regulated peak equals the unregulated one.")
    mm.para(doc, "The adopted transform already shows this turning. Its "
                 "maximum reduction is 67,558 cfs near the 0.2 percent event "
                 "and it declines beyond that, because larger events fill the "
                 "pool sooner. That maximum is a check on itself: 180,550 "
                 "cfs-days divided by 67,558 cfs is 2.7 days, a plausible "
                 "time from flood onset to crest in this basin, so the turn "
                 "in the transform and the reservoir's storage volume "
                 "corroborate one another.")
    mm.figure(doc, "convergence",
              "Figure 8-6. Regulated against unregulated peak at Castle Rock, "
              "with the transform drawn converging on the 1:1 line. The "
              "convergence limb is an ESTIMATE drawn for illustration. It is "
              "not fitted, is used in no result in this memorandum, and does "
              "not affect any value in Sections 8.1 through 8.4 or in the "
              "appendices.")
    mm.para(doc, "The convergence point is drawn rather than fitted because "
                 "fitting does not survive the data. The 48 synthetic members "
                 "cluster between 265,000 and 280,000 cfs unregulated, and "
                 "their reductions there span 1,869 to 99,478 cfs — the "
                 "December 1933 member passes 99 percent of its inflow while "
                 "the December 2015 member, at almost the same magnitude, is "
                 "still holding back 98,000 cfs. Convergence depends on "
                 "hydrograph shape and starting pool, not on peak magnitude "
                 "alone, so a regression through those points is not "
                 "meaningful: fitting all 48 places the crossing at 931,600 "
                 "cfs and fitting the largest twelve places it at 259,600. "
                 "The value shown, 500,000 cfs, is consistent with the "
                 "adopted transform's own declining reduction, which "
                 "extrapolates to zero near 480,000 to 500,000 cfs. It lies "
                 "beyond the 10,000-year unregulated event of 374,643 cfs and "
                 "therefore outside the range of every result presented here. "
                 "Settling it as a number rather than a drawing would require "
                 "scaling additional events through the reservoir model.")


def build_appendix_f(doc, conf):
    mm.h1(doc, "Appendix F. Regulated Frequency Ordinates by Location")
    mm.para(doc, "Regulated peak flow at each of the four locations of "
                 "Section 8, with 95 percent confidence limits. Flows are in "
                 "cubic feet per second. At the gage the unregulated curve is "
                 "shown alongside; downstream, the local contribution is "
                 "shown, being the incremental drainage area scaled off the "
                 "unregulated curve and reduced by the 0.80 timing factor. "
                 "Confidence limits at the three downstream locations are the "
                 "gage's own, translated by the local contribution — see "
                 "Section 8.4.")
    sites = [
        ("Castle Rock gage", 2229, "castle_rock_gage", "Unregulated",
         "cowlitz_unreg_cfs"),
        ("Below Arkansas Creek", 2278, "below_arkansas_creek", "Local",
         "below_arkansas_creek_local_cfs"),
        ("Below Ostrander Creek", 2335, "below_ostrander_creek", "Local",
         "below_ostrander_creek_local_cfs"),
        ("Below Coweeman River", 2476, "below_coweeman_river", "Local",
         "below_coweeman_river_local_cfs"),
    ]
    for n, (label, area, key, second, second_col) in enumerate(sites, start=1):
        rows = []
        for _, r in conf.iterrows():
            rows.append(["%.4f" % r["AEP"], mm.fmt(r[second_col]),
                         mm.fmt(r["%s_cfs" % key]),
                         mm.fmt(r["%s_lower_cfs" % key]),
                         mm.fmt(r["%s_upper_cfs" % key])])
        mm.table(doc, ["AEP", second, "Regulated", "Lower 95%", "Upper 95%"],
                 rows, widths=[1400, 1800, 1800, 1800, 1800],
                 align_right=(0, 1, 2, 3, 4))
        mm.caption(doc, "Table F-%d. %s, %s square miles."
                        % (n, label, "{:,}".format(area)))


def main():
    if not os.path.exists(SRC_DOCX):
        raise SystemExit("Source draft not found: %s" % os.path.abspath(SRC_DOCX))
    shutil.copyfile(SRC_DOCX, OUT_DOCX)
    print("Copied %s -> %s" % (SRC_DOCX, OUT_DOCX))

    mm.FIG.update(NEW_FIGURES)
    conf = pd.read_csv(BELOW_CONF_CSV).sort_values("AEP", ascending=False)
    prism = pd.read_csv(PRISM_CSV)
    tgt = conf.iloc[(conf["AEP"] - 0.001).abs().idxmin()]

    doc = Document(OUT_DOCX)
    body = doc.element.body
    before = len(doc.paragraphs)

    # --- replace Section 8 in place ---
    p8 = find_heading(doc, SECTION8_PREFIX)
    pA = find_heading(doc, APPENDIX_A_PREFIX)
    removed = delete_through(body, p8._p, pA._p)
    print("Removed %d element(s) of the old Section 8 placeholder" % removed)
    els = capture(doc, lambda d: build_section8(d, conf, prism, tgt))
    move_before(pA, els)
    print("Inserted Section 8 (%d elements)" % len(els))

    # --- Appendix F between Appendix E and References ---
    pRef = find_heading(doc, REFERENCES_PREFIX)
    els = capture(doc, lambda d: build_appendix_f(d, conf))
    move_before(pRef, els)
    print("Inserted Appendix F (%d elements)" % len(els))

    doc.save(OUT_DOCX)
    print("Wrote %s  (%d -> %d paragraphs)"
          % (OUT_DOCX, before, len(Document(OUT_DOCX).paragraphs)))
    if mm.MISSING:
        print("MISSING FIGURES: %s" % ", ".join(sorted(set(mm.MISSING))))


if __name__ == "__main__":
    main()
