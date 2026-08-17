#make_memo_reg_unreg.py
# -*- coding: utf-8 -*-
"""
Build the draft regulated flow frequency memorandum.

Structure, styles, page setup and table formatting are taken from the adopted
unregulated memo (TEMPLATE_DOCX) so the two read as one series: the template is
copied, its body is emptied, and the new content is written back using the same
style names and the same direct formatting the original uses.

!! THIS OVERWRITES OUT_DOCX !!
A memo gets hand-edited in Word -- tracked changes, review comments, figure
placement. Re-running this script throws all of that away, the same way
#Create_Synthetic_Ensembles.py throws away the hand-chopped members. Once the
first round of review starts, treat the .docx as the source of record and this
script as history. If a structural change is needed after that, make it in
Word, not here.

Numbers are pulled from the analysis outputs where a file exists, so a re-run
before review picks up the current results rather than restating stale ones.
Anything not in a CSV is stated inline and marked in REVIEW_NOTES.
"""

import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

import copy
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Use the CLEAN unregulated memo, not the _DQC_comments copy. That copy carries
# word/comments.xml and its four companion parts; emptying the body strips the
# comment anchors out of document.xml but leaves those parts behind, and the
# orphaned result will not open in LibreOffice. It also has no business
# carrying another document's review comments into a new draft.
TEMPLATE_DOCX = r"../../CAS_Unreg_FF/docs/MEMO_CAS_Unreg_FF_04Aug2026.docx"
OUT_DOCX = r"MEMO_CAS_Reg_Unreg_DRAFT.docx"

REG_FREQ_CSV = r"../output/regulated_frequency_inferred.csv"
ADJUSTED_CSV = r"../output/adjusted_peaks.csv"
CRITDUR_CSV = r"../output/critical_duration_adjusted_fits.csv"
SYNTH_CSV = r"../output/synthetic_results.csv"

FIG = {
    "basin": r"../../CAS_Unreg_FF/docs/figures/basin_map.png",
    "adjusted": r"../output/diagnostics/adjusted_peaks.png",
    "critdur": r"../output/diagnostics/critical_duration_adjusted_scatter.png",
    "synth_events": r"../output/diagnostics/ensemble_synthetic_events.png",
    "scatter": r"../output/diagnostics/unreg_reg_scatter_loglog.png",
    "final": r"../output/diagnostics/unreg_reg_final_uncertainty.png",
    "cmp2009": r"../output/diagnostics/unreg_reg_2009_comparison.png",
}
FIG_WIDTH_IN = 6.5

# Formatting lifted from the template, in half-points / twentieths of a point.
SZ_BODY = None          # inherit Normal
SZ_TABLE = 19           # 9.5 pt
SZ_CAPTION = 19
SZ_H1, SZ_H2 = 26, 23
SZ_TITLE = 28
HDR_FILL = "D9E2EC"
BORDER_OUTER, BORDER_INNER = "7A869A", "B7C0CC"
TABLE_WIDTH_DXA = 9360

# ----------------------------------------------------------------------------


def fmt(value, digits=0):
    """Thousands-separated number, or an em dash when it is not available."""
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "—"
    return format(round(float(value), digits), ",.%df" % digits)


def clear_body(doc):
    """Empty the body but keep the final sectPr (page size, margins, footer)."""
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect:
            body.remove(child)
    return doc


def assert_no_comment_parts(path):
    """Refuse a template that carries comments.

    Emptying the body removes the anchors but not word/comments.xml and its
    companions, and the orphaned parts make the output unopenable. Checked
    rather than assumed, because the two memo copies differ only by a suffix
    and picking the wrong one produces a file that looks fine until someone
    tries to open it.
    """
    import zipfile
    with zipfile.ZipFile(path) as archive:
        bad = [n for n in archive.namelist() if "comments" in n.lower()]
    if bad:
        raise SystemExit(
            "TEMPLATE_DOCX carries comment parts and cannot be used as a "
            "template:\n   %s\n   %s\nUse the clean memo instead."
            % (path, ", ".join(bad)))


def para(doc, text="", size=None, bold=False, italic=False, align=None,
         style=None, before=None, after=160, line=276):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = Pt(before / 20.0)
    if after is not None:
        pf.space_after = Pt(after / 20.0)
    if line:
        pf.line_spacing = line / 240.0
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size / 2.0)
    return p


def h1(doc, text):
    return para(doc, text, size=SZ_H1, bold=True, style="Heading 1",
                before=300, after=140, line=None)


def h2(doc, text):
    return para(doc, text, size=SZ_H2, bold=True, style="Heading 2",
                before=300, after=140, line=None)


def caption(doc, text):
    return para(doc, text, size=SZ_CAPTION, italic=True,
                align=WD_ALIGN_PARAGRAPH.CENTER, before=80, after=240,
                line=None)


def figure(doc, key, cap, width_in=FIG_WIDTH_IN):
    """Insert a figure, or a visible placeholder if the file is missing.

    A missing figure must not be silent -- a memo that quietly drops one is
    worse than one that says a figure is not here yet.
    """
    path = FIG.get(key)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    if path and os.path.isfile(path):
        p.add_run().add_picture(path, width=Inches(width_in))
    else:
        run = p.add_run("[FIGURE NOT FOUND: %s]" % (path or key))
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2D)
        MISSING.append(path or key)
    caption(doc, cap)


def _borders(tbl):
    pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, size, color in (("top", 4, BORDER_OUTER), ("left", 4, BORDER_OUTER),
                              ("bottom", 4, BORDER_OUTER), ("right", 4, BORDER_OUTER),
                              ("insideH", 2, BORDER_INNER), ("insideV", 2, BORDER_INNER)):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    pr.append(borders)


def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def table(doc, header, rows, widths=None, align_right=()):
    """A table formatted the way the template formats its tables."""
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = doc.styles["Normal Table"]
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _borders(tbl)

    if widths is None:
        widths = [TABLE_WIDTH_DXA // len(header)] * len(header)

    def write(cell, text, bold, idx):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if bold:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif idx in align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(SZ_TABLE / 2.0)

    for i, text in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(widths[i] / 1440.0)
        _shade(cell, HDR_FILL)
        write(cell, text, True, i)
    for row in rows:
        cells = tbl.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Inches(widths[i] / 1440.0)
            write(cells[i], text, False, i)
    return tbl


def load(path, what):
    if not os.path.isfile(path):
        REVIEW_NOTES.append("%s not found (%s) -- numbers from it are blank."
                            % (what, path))
        return None
    return pd.read_csv(path)


MISSING = []
REVIEW_NOTES = []


def main():
    reg = load(REG_FREQ_CSV, "Regulated frequency table")
    adj = load(ADJUSTED_CSV, "Adjusted peak record")
    crit = load(CRITDUR_CSV, "Critical duration fits")
    synth = load(SYNTH_CSV, "Synthetic results")

    def at(aep, col):
        """Value at the AEP nearest `aep`.

        idxmin, not argsort()[:1] -- the latter returns a Series whose values
        are positions but whose index is the original index, and .iloc on it
        silently picked the LAST row every time, so every number quoted from
        the frequency table was the 0.01 percent ordinate.
        """
        if reg is None or col not in reg.columns:
            return None
        i = (reg["AEP"] - aep).abs().idxmin()
        value = reg.loc[i, col]
        return None if pd.isna(value) else float(value)

    assert_no_comment_parts(TEMPLATE_DOCX)
    doc = clear_body(Document(TEMPLATE_DOCX))

    # ---------------------------------------------------------------- header
    para(doc, "MEMORANDUM FOR RECORD", size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=60, line=None)
    para(doc, "U.S. Army Corps of Engineers, Portland District", size=21,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=240, line=None)
    table(doc, ["Item", "Description"], [
        ["Subject", "Regulated Flow Frequency, Cowlitz River at Castle Rock, "
                    "Washington (USGS 14243000)"],
        ["From", "Hydrology and Hydraulics Branch, Portland District"],
        ["Date", "August 2026"],
        ["Status", "DRAFT — Section 14 pending"],
    ], widths=[2400, 6960])
    para(doc, "", after=120)

    # ------------------------------------------------------------- 1 Purpose
    h1(doc, "1. Purpose")
    para(doc, "This memorandum documents the development of regulated flow "
              "frequency curves for the Cowlitz River at Castle Rock, "
              "Washington. The regulated curves are obtained by transforming "
              "the adopted unregulated curves through a relationship between "
              "unregulated and regulated peak flow derived from reservoir "
              "simulation, rather than by fitting a distribution to the "
              "observed regulated record.")
    para(doc, "The unregulated frequency analysis that supplies the input to "
              "this study is documented separately in the memorandum for "
              "record on unregulated flow frequency at Castle Rock (Portland "
              "District, August 2026). This memorandum begins where that one "
              "ends: with an adopted unregulated frequency curve for the "
              "instantaneous peak, and takes it through the projects to a "
              "regulated curve with uncertainty. The regulated curves support "
              "levee fragility assessment and floodplain mapping in the lower "
              "Cowlitz valley.")
    para(doc, "Regulated peak flows cannot be fit directly. Two independent "
              "reasons apply. First, the observed regulated record is not a "
              "homogeneous sample: Mossyrock began impounding in December "
              "1968, and operating rules and starting pool practice have "
              "varied since, so the regulated peaks in the record were not "
              "all produced by the same system. Second, regulated peaks do "
              "not follow an analytical distribution. Operating rules put "
              "hard breaks in the relationship between inflow and outflow "
              "— a reservoir holds small events almost entirely, loses "
              "ground through the middle of its range, and passes the largest "
              "events through once storage is exhausted — and no "
              "log-Pearson Type III curve reproduces that shape. The standard "
              "approach, and the one used here, is to establish the frequency "
              "of the natural inflow and carry it through the projects.")

    # ------------------------------------------------- 2 Basin and operation
    h1(doc, "2. Basin Description and Project Operation")
    para(doc, "The Cowlitz River drains the southern Washington Cascades and "
              "enters the Columbia River near Longview. The Castle Rock gage "
              "sits in the lower valley, downstream of the Toutle River "
              "confluence and downstream of all the basin's storage projects. "
              "The total basin area above Castle Rock is 2,238 square miles, "
              "of which 1,170 square miles of headwaters are regulated by "
              "Mossyrock Dam.")
    para(doc, "Three non-Federal hydroelectric projects lie on the mainstem, "
              "operated by Tacoma Power as the Cowlitz Hydroelectric Project. "
              "Mossyrock Dam, completed in 1968, impounds Riffe Lake and "
              "provides the only usable flood storage in the basin. Mayfield "
              "Dam, completed in 1963, is a re-regulating dam that smooths "
              "power peaking releases from Mossyrock and provides no "
              "meaningful flood regulation of its own. Cowlitz Falls Dam, "
              "above Mossyrock, is run-of-river.")
    para(doc, "Two features of the basin govern how much the projects can do, "
              "and both are visible in the results that follow. The regulated "
              "fraction is a little over half the drainage area, so roughly "
              "half the flow arriving at Castle Rock in a large event is "
              "uncontrolled. The Toutle River, the largest tributary, enters "
              "below both dams and is entirely unregulated; the Tilton River "
              "enters Mayfield Lake and is caught by the projects. At the "
              "largest events the local and Toutle contribution is a median "
              "48 percent of unregulated flow at Castle Rock, which sets a "
              "floor on the regulated peak that no operation can reach below.")
    figure(doc, "basin", "Figure 2-1. Cowlitz River basin, showing the Castle "
                         "Rock gage, the Mossyrock and Mayfield projects, and "
                         "the Toutle and Tilton tributaries.")

    h2(doc, "2.1 Previous Studies")
    para(doc, "The previously adopted regulated frequency curve at Castle "
              "Rock comes from the 2009 Cowlitz Hydrology Restudy. That study "
              "predates Bulletin 17C, does not include the most recent 17 "
              "years of record, and derived its regulated curve by a "
              "different route. Section 11 compares the two.")

    # ------------------------------------------------------------ 3 Data used
    h1(doc, "3. Data Used")
    para(doc, "Table 3-1 lists the datasets and models used and the purpose "
              "each one served. The observed records are the same ones "
              "assembled for the unregulated study and are documented there; "
              "the additions specific to this study are the reservoir "
              "simulation model and the simulation outputs it produced.")
    table(doc, ["Dataset or model", "Period used", "Purpose"], [
        ["Adopted unregulated frequency curve, Castle Rock",
         "WY 1927–2026",
         "Input to the transform; supplies the AEP of every regulated ordinate"],
        ["HEC-ResSim model, Mossyrock and Mayfield",
         "Current (2014 WCM) rules",
         "Routes inflow hydrographs through the projects"],
        ["USGS 14243000 annual peak record", "WY 1927–2026",
         "Observed regulated peaks; basis of the adjusted record"],
        ["Mossyrock pool elevation, daily (USGS)", "Oct 1973–present",
         "Observed starting pool for the Obs_RC simulation"],
        ["Reservoir inflow and Castle Rock local flow",
         "WY 1929–2026",
         "Simulation inflows, back-calculated and volume-corrected"],
        ["ResSim period-of-record simulation, rule-curve start (WCM_RC)",
         "98 water years", "Regulated peak under current rules from a "
                           "rule-curve starting pool"],
        ["ResSim period-of-record simulation, observed start (Obs_RC)",
         "53 water years", "Regulated peak under current rules from the "
                           "observed starting pool"],
        ["ResSim synthetic ensemble simulation", "48 members",
         "Regulated response above the observed range"],
    ], widths=[3400, 1900, 4060])
    caption(doc, "Table 3-1. Datasets and models used, and their purpose.")

    # -------------------------------------------------------------- 4 Methods
    h1(doc, "4. Methods")
    para(doc, "The regulated curve is built in four steps, each documented in "
              "its own section below.")
    para(doc, "First, an adjusted regulated peak record is assembled from the "
              "observed record, placing every historical peak on a consistent "
              "starting-pool basis so that peaks from different decades can "
              "be compared with one another (Section 5). Second, those "
              "adjusted peaks are paired with the unregulated peak for the "
              "same event, giving a set of observed (unregulated, regulated) "
              "pairs (Section 6). Third, because only one water year in the "
              "record exceeds the unregulated 100-year peak, synthetic flood "
              "ensembles are constructed and routed to populate the range "
              "above the record, where the answer is needed and the record is "
              "silent (Section 7). Fourth, a relationship is drawn through "
              "the combined set of pairs and applied to the unregulated "
              "frequency curve, transferring the AEP of each unregulated "
              "ordinate to the regulated flow it produces (Section 8).")
    para(doc, "The transform is applied peak-to-peak. The instantaneous peak "
              "and the 1-day duration are effectively tied as the critical "
              "duration for the regulated response (Section 6), and a "
              "peak-to-peak transform is the simpler of the two to explain "
              "and to audit.")

    # ------------------------------------------------- 5 Adjusted peak record
    h1(doc, "5. Adjusted Regulated Peak Record")
    para(doc, "The observed regulated peaks are the direct evidence of what "
              "the projects do to a flood, but they are not directly "
              "comparable to one another. A flood arriving on a full pool is "
              "attenuated less than the same flood arriving on a drawn-down "
              "pool, so a peak from a year that happened to start low "
              "understates what the same storm would do under the operation "
              "being evaluated. The record has to be placed on a consistent "
              "starting-pool basis before it can be used.")
    para(doc, "Two period-of-record ResSim simulations were run for this "
              "purpose. Both use the same observed hydrology and the same "
              "current operating rules, and differ only in the reservoir's "
              "starting pool: WCM_RC starts each event at the Water Control "
              "Manual rule curve, and Obs_RC starts it at the observed "
              "elevation at event onset. Because everything else is held "
              "fixed, the difference between them isolates the effect of the "
              "starting pool alone. Obs_RC is limited to water year 1974 "
              "onward because the observed daily Mossyrock elevation record "
              "begins in October 1973.")
    para(doc, "Where the rule-curve start produces a higher regulated peak "
              "than the observed start, the historical operation benefited "
              "from storage the rule curve does not credit. That benefit is "
              "removed from the observed peak:")
    para(doc, "adjusted peak = observed peak + (WCM_RC peak − Obs_RC "
              "peak),  where WCM_RC exceeds Obs_RC",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=160)
    para(doc, "The adjustment is one-sided by design: the observed peak is "
              "only ever increased, never reduced. A negative difference "
              "means the observed pool started above the rule curve, so the "
              "historical operation was already at least as conservative as "
              "the rule curve requires and there is nothing to remove. Only "
              "water year 1976 falls in that category.")
    para(doc, "Peaks are matched by event rather than by water year. Each "
              "simulation's peak is taken within three days of the observed "
              "peak date, not as that run's annual maximum. This matters "
              "because a low starting pool attenuates the first storm of a "
              "window more, which can promote a later storm to the annual "
              "maximum; differencing annual maxima would then subtract two "
              "different storms. Event matching yields 40 usable years "
              "against 29 for annual maxima.")

    if adj is not None:
        n_shared = len(adj)
        n_adjusted = int(adj["adjusted"].sum()) if "adjusted" in adj else None
        deltas = adj.loc[adj["adjusted"] == True, "delta_wcm_minus_obs"] \
            if "adjusted" in adj else pd.Series(dtype=float)
        codes = adj["screen_code"].value_counts() if "screen_code" in adj else {}
        # Count each failing year once. "different_event" and "reg_over_unreg"
        # are exclusive of "both", so adding "both" into each total reports
        # more failures than there are years.
        n_timing = int(codes.get("different_event", 0))
        n_physical = int(codes.get("reg_over_unreg", 0))
        n_both = int(codes.get("both", 0))
        n_failed = n_timing + n_physical + n_both
        para(doc, "Of %d shared water years (%d through %d), %d received an "
                  "adjustment, with a median of %s cubic feet per second and a "
                  "range of %s to %s. Screening removed %d years from the fit: "
                  "%d failed a timing screen, where the observed and simulated "
                  "peaks are not the same storm; %d failed a physical screen in "
                  "which the regulated peak exceeded the unregulated peak at or "
                  "above 60,000 cubic feet per second, which a reservoir cannot "
                  "produce; and %d failed both. Screened years are carried "
                  "through the record unadjusted rather than deleted, so the "
                  "reason is recorded against each one in Appendix A."
              % (n_shared, int(adj["WY"].min()), int(adj["WY"].max()),
                 n_adjusted, fmt(deltas.median()), fmt(deltas.min()),
                 fmt(deltas.max()), n_failed, n_timing, n_physical, n_both))
    para(doc, "Water year 1980 is excluded explicitly. The observed 97,000 "
              "cubic feet per second peak on 18 May 1980 is the Mount St. "
              "Helens eruption, a debris-flow event with no meteorological "
              "analogue; both simulations peak five months earlier on an "
              "ordinary winter storm. Whether that event belongs in a flood "
              "frequency analysis at all is a separate question this study "
              "does not decide.")
    figure(doc, "adjusted", "Figure 5-1. Observed, WCM_RC and Obs_RC regulated "
                            "peaks by water year, and the resulting adjustment.")

    # --------------------------------------------------- 6 Critical duration
    h1(doc, "6. Critical Duration")
    para(doc, "The duration at which the unregulated and regulated records "
              "correspond most closely identifies the averaging window over "
              "which the projects actually operate, and therefore the "
              "duration the transform should be built on. Adjusted regulated "
              "values were regressed on unregulated values at the "
              "instantaneous peak and at the 1-, 3- and 5-day durations.")
    if crit is not None:
        rows = []
        for _, r in crit.iterrows():
            rows.append([str(r["duration"]), "%d" % r["n"],
                         "%.3f" % r["r2"], "%.3f" % r["log_r2"],
                         "%.3f" % r["log_exponent_b"],
                         "%.2f" % r["median_reg_over_unreg"]])
        table(doc, ["Duration", "n", "R-squared", "Log R-squared",
                    "Log slope", "Median reg / unreg"], rows,
              widths=[2000, 800, 1600, 1700, 1500, 1760],
              align_right=(1, 2, 3, 4, 5))
        caption(doc, "Table 6-1. Adjusted regulated against unregulated flow "
                     "by duration, water years 1974 through 2024.")
    para(doc, "The instantaneous peak and the 1-day duration are not "
              "distinguishable on this evidence, and both are stronger than "
              "the longer durations. The peak-to-peak relation was adopted. "
              "The median ratio column shows the physical reason the shorter "
              "durations correlate better: at the peak the projects reduce "
              "flow to about 78 percent of unregulated, while at 5 days the "
              "ratio exceeds one, because the water held out of a peak is "
              "released afterward and reappears in the longer averaging "
              "window. The projects move flood volume in time; they do not "
              "remove it.")
    figure(doc, "critdur", "Figure 6-1. Adjusted regulated against "
                           "unregulated flow, by duration.")

    # ----------------------------------------------------- 7 Synthetic events
    h1(doc, "7. Synthetic Flood Ensembles")
    para(doc, "Only one water year in the record exceeds the unregulated "
              "100-year peak, so the upper end of the unregulated-regulated "
              "relationship — the part that sets the 100-, 250- and "
              "500-year regulated flows — is essentially unconstrained "
              "by observed events. Synthetic floods were constructed and "
              "routed to populate that range.")
    para(doc, "The design is a factorial over the two things that could "
              "matter. Magnitude is varied over four targets taken from the "
              "unregulated frequency curve: the 100-, 250- and 500-year peaks, "
              "and a fourth target 20 percent above the 500-year so that the "
              "drawn relationship is not extrapolating past its final point "
              "exactly where the answer is needed. Shape is varied by scaling "
              "twelve different observed storms, selected from the record by "
              "unregulated peak and spanning the full range of observed "
              "attenuation behavior, from storms the projects absorbed to one "
              "that passed straight through. Shape is varied rather than "
              "assumed because attenuation among large observed events is not "
              "predictable from shape, magnitude, antecedent flow or timing.")
    para(doc, "Each source storm is scaled to hit its target unregulated peak "
              "exactly and its target 5-day volume by iteration, using a "
              "multiplier that is a function of flow rather than of time. A "
              "flow-based multiplier cannot reorder the hydrograph, so the "
              "peak stays the peak; a time-based taper can scale a shoulder "
              "harder than the peak and invent a new maximum. Both the "
              "Mossyrock inflow and the Castle Rock local are scaled by the "
              "same factor, so the observed coincidence between the "
              "controlled and uncontrolled fractions is preserved.")
    para(doc, "Two members required manual editing. Matching a peak and a "
              "5-day volume on a sharp storm forces the shoulders to stretch "
              "harder than the peak itself, which turned a small rise already "
              "present in the December 1977 and November 1986 hydrographs "
              "into what read as a separate flood ahead of the main peak. "
              "Both were removed by hand before the ensemble was routed. The "
              "effect on the result is confined to four members, whose "
              "regulated peaks fell by 1.4 to 5.4 percent; the remaining 44 "
              "are unchanged.")
    if synth is not None and "reg_peak" in synth:
        # Ratio computed the same way the transform pairs are: routed
        # regulated over routed unregulated. The atten_ratio column in the
        # results CSV is taken against the unrouted build target and reads a
        # couple of percent lower, which would not match Section 8.
        synth = synth.copy()
        synth["ratio"] = synth["reg_peak"] / synth["unreg_peak_routed_cfs"]
        by_event = synth.groupby("event")["ratio"]
        lo_evt, hi_evt = by_event.median().idxmin(), by_event.median().idxmax()
        para(doc, "The %d routed members span regulated peaks from %s to %s "
                  "cubic feet per second and attenuation ratios, regulated "
                  "over unregulated peak, from %.2f to %.2f. The spread at a "
                  "single magnitude is a shape effect "
                  "rather than scatter: the %s storm is absorbed most "
                  "strongly, while %s exhausts storage and passes through near "
                  "the 1:1 line, and the two land far apart at the same "
                  "unregulated peak. That spread is the physical basis for the "
                  "uncertainty band in Section 10."
              % (len(synth), fmt(synth["reg_peak"].min()),
                 fmt(synth["reg_peak"].max()),
                 synth["ratio"].min(), synth["ratio"].max(), lo_evt, hi_evt))
    figure(doc, "synth_events", "Figure 7-1. Source storms and the scaled "
                                "family built from each. The shaded band is "
                                "the window over which the 5-day volume is "
                                "matched.")

    # ------------------------------------------------------- 8 The transform
    h1(doc, "8. The Unregulated-Regulated Transform")
    para(doc, "The adjusted historical pairs and the routed synthetic members "
              "together give the set of (unregulated, regulated) pairs the "
              "transform is drawn through. The historical pairs carry the "
              "observed range; the synthetics carry the range above it.")
    para(doc, "A single power law was considered and rejected as the adopted "
              "form. A power law is a straight line in log-log space, and a "
              "straight line is the wrong shape for this relationship: the "
              "projects hold small events almost entirely, lose ground "
              "through the middle of the range, and flatten toward the "
              "unregulated flow at the top as storage is exhausted. Forcing "
              "one slope through all of that places the line above the data "
              "in one range and below it in another, and the error lands at "
              "the upper end, where it matters most.")
    para(doc, "The adopted transform is a locally weighted regression through "
              "the scatter in log-log space, using tricube weights over the "
              "nearest 65 percent of the sample, so the slope at any point "
              "comes from nearby data only and the line is free to bend where "
              "the data bends. It is constrained to increase monotonically, "
              "and clipped at the 1:1 line at large flows, since a regulated "
              "peak above its own unregulated peak is not physical there. The "
              "single power law is retained as a reference line on the "
              "diagnostic figures.")
    para(doc, "The WCM_RC simulated pairs are shown on the diagnostic figures "
              "but are deliberately not fitted. They are simulated on both "
              "axes and carry no observed anchor: the adjustment that "
              "corrects the historical record is built from the difference "
              "between the two simulations, and where no Obs_RC data exists "
              "there is no correction to apply. They are evidence about the "
              "shape of the relationship, not about its position.")
    figure(doc, "scatter", "Figure 8-1. Unregulated against regulated peak at "
                           "Castle Rock, log-log, with the adopted transform "
                           "and the single power law for reference.")

    # ------------------------------------------------- 9 The regulated curve
    h1(doc, "9. Regulated Frequency Curve")
    para(doc, "The regulated curve is obtained by reading each ordinate of "
              "the adopted unregulated curve through the transform. The "
              "annual exceedance probability is inherited from the "
              "unregulated side: the regulated flow plotted at a given AEP is "
              "the flow the projects produce from the unregulated flow of "
              "that AEP. No distribution is fitted to regulated flow at any "
              "point in this procedure.")
    if reg is not None:
        rows = []
        for aep, label in [(0.5, "2-year"), (0.1, "10-year"), (0.02, "50-year"),
                           (0.01, "100-year"), (0.005, "200-year"),
                           (0.002, "500-year")]:
            rows.append([
                label, "%.3f" % aep,
                fmt(at(aep, "unreg_expected_cfs")),
                fmt(at(aep, "reg_inferred_cfs")),
                "%.0f%%" % (at(aep, "reduction_pct") or 0),
                "%s – %s" % (fmt(at(aep, "reg_lower_90pct_cfs")),
                                  fmt(at(aep, "reg_upper_90pct_cfs")))])
        table(doc, ["Event", "AEP", "Unregulated (cfs)", "Regulated (cfs)",
                    "Reduction", "Regulated, 90% band (cfs)"], rows,
              widths=[1200, 900, 1800, 1700, 1200, 2560],
              align_right=(1, 2, 3, 4, 5))
        caption(doc, "Table 9-1. Adopted regulated frequency curve at Castle "
                     "Rock, selected ordinates.")
        para(doc, "The reduction the projects achieve is not monotonic in "
                  "magnitude. It peaks at about %.0f percent near the 1 to 2 "
                  "percent annual chance event and falls away on both sides: "
                  "below, because small events are already largely contained "
                  "and the unregulated flow is itself small; above, because "
                  "storage is finite and the largest floods increasingly pass "
                  "through. At the 0.2 percent chance event the reduction is "
                  "down to about %.0f percent. This shape, not any property of "
                  "a fitted distribution, is what the transform exists to "
                  "capture."
              % (reg["reduction_pct"].max(),
                 at(0.002, "reduction_pct") or 0))
    figure(doc, "final", "Figure 9-1. Adopted regulated and unregulated "
                         "frequency curves at Castle Rock, with the 90 percent "
                         "uncertainty band.")

    # ------------------------------------------------------- 10 Uncertainty
    h1(doc, "10. Uncertainty")
    para(doc, "The regulated flow at a given AEP is uncertain for two "
              "independent reasons, and the band carries both.")
    para(doc, "The first is flow frequency. The unregulated quantile being "
              "read through the transform is itself an estimate from a "
              "roughly 100-year record, and HEC-SSP reports its uncertainty "
              "as confidence limits about the computed curve. The second is "
              "the transform itself. At one unregulated magnitude the "
              "regulated peak still depends on the shape of the flood, which "
              "is what the scatter of the pairs about the transform measures.")
    para(doc, "The two are combined in log space, but the frequency term must "
              "be carried through the transform before they are added. A "
              "given proportional error in unregulated flow does not produce "
              "the same proportional error in regulated flow: it is scaled by "
              "the local slope of the transform in log-log space. That slope "
              "runs about 0.57 through the middle of the curve, where the "
              "projects absorb most of an increase, and rises to about 1.5 at "
              "the top, where the transform bends toward pass-through and "
              "regulated flow is catching unregulated flow up. Omitting that "
              "term would overstate the band through the middle by nearly a "
              "factor of two and understate it at the top.")
    para(doc, "The confidence limits reported by HEC-SSP are asymmetric, "
              "being roughly 1.5 times wider on the upper side than the lower "
              "at the 500-year. That asymmetry is preserved by combining the "
              "two sides separately rather than averaging them, which remains "
              "a closed-form calculation. The resulting band is at most 1.4 "
              "times more lopsided on one side than the other, which is well "
              "within the range a two-piece lognormal represents adequately; "
              "no Monte Carlo sampling was required.")
    para(doc, "The scatter of the transform is not constant along the curve "
              "and is not treated as constant. Measured locally it runs about "
              "0.05 in log units below 80,000 cubic feet per second and about "
              "0.08 above 150,000. This is physical: a small flood is simply "
              "held, while whether a large one exhausts storage depends on "
              "its shape. A single pooled value would be about 25 percent too "
              "wide at the median and 15 percent too narrow at the top.")
    para(doc, "One constraint is imposed on the band. Above 60,000 cubic feet "
              "per second the upper bound of the regulated curve is held at "
              "or below the upper bound of the unregulated curve, because a "
              "regulated flood cannot exceed the unregulated flood it was "
              "routed from. Below that threshold the constraint is not "
              "applied: minimum release requirements and refill drawdown can "
              "legitimately put more water in the river than natural "
              "conditions would. The constraint binds only at the three "
              "smallest AEPs on the curve.")

    # ---------------------------------------------------- 12 2009 comparison
    h1(doc, "11. Comparison with the 2009 Restudy")
    if reg is not None and "reg_vs_2009_pct" in reg:
        d100 = at(0.01, "reg_vs_2009_pct") or 0.0
        d500 = at(0.002, "reg_vs_2009_pct") or 0.0
        d1000 = at(0.001, "reg_vs_2009_pct") or 0.0
        mid = abs(reg.loc[reg["AEP"].between(0.05, 0.5),
                          "reg_vs_2009_pct"]).max()
        para(doc, "Figure 11-1 places the adopted regulated curve against the "
                  "2009 Cowlitz Hydrology Restudy. Through the range supported "
                  "by observed events the two agree within about %.0f percent. "
                  "At the 1 percent annual chance event the 2026 curve is %.0f "
                  "percent %s the 2009 value, and at the 0.2 percent chance "
                  "event %.0f percent %s it. Beyond that the 2026 curve falls "
                  "well below, by %.0f percent at the 0.1 percent chance event."
              % (mid, abs(d100), "above" if d100 >= 0 else "below",
                 abs(d500), "above" if d500 >= 0 else "below", abs(d1000)))
    para(doc, "The divergence at the extreme tail is driven by the shape of "
              "the 2009 curve, which turns sharply upward above the 0.2 "
              "percent annual chance event, rising from 156,000 to 390,000 "
              "cubic feet per second between the 0.2 and 0.01 percent chance "
              "events. Any comparison in that range is between two "
              "extrapolations rather than between two records, and should be "
              "read as such.")
    figure(doc, "cmp2009", "Figure 11-1. Adopted 2026 regulated curve against "
                           "the 2009 Cowlitz Hydrology Restudy.")

    # -------------------------------------------------------- 11 Limitations
    h1(doc, "12. Limitations")
    para(doc, "The transform is supported by data over the unregulated range "
              "of roughly 23,000 to 280,000 cubic feet per second. The upper "
              "part of that range is supported entirely by synthetic events, "
              "not by observed floods, and the synthetic events are observed "
              "storms scaled up. Hydrograph shape is not independent of size "
              "— a large storm is large because it was widespread and "
              "sustained — so scaling a moderate storm to a rare "
              "magnitude routes a shape that storm never produced. Six of the "
              "twelve source storms must be scaled past a factor of two to "
              "reach the largest target, and those members should be read as "
              "the weaker evidence in the set.")
    para(doc, "The adjusted record carries its own caveats. Thirteen of the "
              "adjustments land within a narrow band near 14,000 cubic feet "
              "per second, which is most likely a fixed release constraint "
              "binding in one simulation and not the other rather than a pure "
              "starting-pool effect. Seven adjustments exceed 40 percent of "
              "the observed peak; all are moderate-peak years where a large "
              "absolute adjustment is a large fraction, and they influence the "
              "lower end of the transform more than the upper.")
    para(doc, "The adjusted record is deliberately not used for an analytical "
              "frequency fit, and should not be. It exists to establish the "
              "unregulated-regulated relationship at large events, which is "
              "why the screening losses are acceptable: the excluded years are "
              "dominated by low peaks that do not influence that relationship.")
    para(doc, "Finally, the transform is a centre-of-mass line through a "
              "scatter, not a physical model of the projects. It reproduces "
              "what the simulations did on the events tested. It carries no "
              "information about operations outside the current rule set, and "
              "it should be redrawn if those rules change.")

    # ------------------------------------------- 13 Downstream extension TBD
    h1(doc, "13. Extension to the Columbia River Confluence")
    para(doc, "[TO BE COMPLETED]", bold=True)
    para(doc, "The regulated frequency curve developed above applies at "
              "Castle Rock. Extending it downstream to the Columbia River "
              "confluence requires a coincident frequency analysis combining "
              "the regulated Cowlitz flow with Columbia River stage, "
              "supported by regional flood frequency estimates for the "
              "intervening ungaged area. That analysis is in progress and "
              "will be documented in this section.")
    para(doc, "This section is expected to cover the coincident frequency "
              "method and the assumed dependence between Cowlitz and Columbia "
              "flooding; the regional regression estimates used for the "
              "ungaged local area below Castle Rock; the resulting regulated "
              "frequency estimates at the downstream locations required for "
              "levee analysis; and how the uncertainty developed in Section "
              "10 is carried through that combination.")

    # ------------------------------------------------------------ Appendices
    h1(doc, "Appendix A. Adjusted Regulated Peak Record by Water Year")
    para(doc, "Observed, simulated and adjusted regulated peaks for every "
              "shared water year, with the screening result and the reason "
              "recorded for each. Flows are in cubic feet per second.")
    if adj is not None:
        cols = [c for c in ["WY", "usgs", "wcm", "obs", "delta_wcm_minus_obs",
                            "adjusted_peak", "screen_code"] if c in adj.columns]
        rows = []
        for _, r in adj.sort_values("WY").iterrows():
            rows.append([("%d" % r["WY"]) if c == "WY"
                         else (str(r[c]) if c == "screen_code" else fmt(r[c]))
                         for c in cols])
        table(doc, ["WY", "Observed", "WCM_RC", "Obs_RC", "Difference",
                    "Adjusted", "Screen"][:len(cols)], rows,
              widths=[900, 1400, 1400, 1400, 1400, 1400, 1460][:len(cols)],
              align_right=tuple(range(1, len(cols) - 1)))
        caption(doc, "Table A-1. Adjusted regulated peak record, water years "
                     "%d through %d." % (adj["WY"].min(), adj["WY"].max()))

    h1(doc, "Appendix B. Regulated Frequency Curve Ordinates")
    para(doc, "The adopted regulated curve with its uncertainty band, and the "
              "terms that produced the band, for every ordinate. The local "
              "slope column is the slope of the transform in log-log space at "
              "that ordinate. Flows are in cubic feet per second.")
    if reg is not None:
        rows = []
        for _, r in reg.iterrows():
            rows.append(["%.4f" % r["AEP"], fmt(r["unreg_expected_cfs"]),
                         fmt(r["reg_inferred_cfs"]),
                         fmt(r.get("reg_lower_90pct_cfs")),
                         fmt(r.get("reg_upper_90pct_cfs")),
                         "%.2f" % r.get("transform_slope_b", float("nan")),
                         "%.1f" % r["reduction_pct"],
                         "yes" if r.get("extrapolated") else "no"])
        table(doc, ["AEP", "Unregulated", "Regulated", "Lower 90%",
                    "Upper 90%", "Local slope", "Reduction %", "Extrap."],
              rows, widths=[1000, 1400, 1300, 1300, 1300, 1150, 1150, 760],
              align_right=(0, 1, 2, 3, 4, 5, 6))
        caption(doc, "Table B-1. Regulated frequency curve ordinates and "
                     "uncertainty terms.")

    h1(doc, "References")
    for ref in [
        "Lind, G.D., Lamontagne, J.R., and Stonewall, A.J., 2020, Regional "
        "skew for selected flood durations for the Columbia River basin: U.S. "
        "Geological Survey Scientific Investigations Report.",
        "Mastin, M.C., Konrad, C.P., Veilleux, A.G., and Tecca, A.E., 2016, "
        "Magnitude, frequency, and trends of floods at gaged and ungaged "
        "sites in Washington, based on data through water year 2014: U.S. "
        "Geological Survey Scientific Investigations Report 2016-5118.",
        "Tacoma Power, 2014, Water Control Manual, Cowlitz Hydroelectric "
        "Project, Mossyrock and Mayfield Dams.",
        "U.S. Army Corps of Engineers, 2009, Cowlitz River Hydrology "
        "Restudy: Portland District.",
        "U.S. Army Corps of Engineers, 2026, Memorandum for Record: "
        "Unregulated Flow Frequency, Cowlitz River at Castle Rock, "
        "Washington: Portland District.",
        "U.S. Army Corps of Engineers, Hydrologic Engineering Center, HEC-ResSim "
        "Reservoir System Simulation, and HEC-SSP Statistical Software "
        "Package.",
        "U.S. Geological Survey, National Water Information System: "
        "streamflow and peak-flow records for station 14243000, Cowlitz River "
        "at Castle Rock, Washington.",
        "England, J.F., Jr., Cohn, T.A., Faber, B.A., Stedinger, J.R., "
        "Thomas, W.O., Jr., Veilleux, A.G., Kiang, J.E., and Mason, R.R., Jr., "
        "2018, Guidelines for determining flood flow frequency—Bulletin "
        "17C: U.S. Geological Survey Techniques and Methods, book 4, chap. B5.",
    ]:
        para(doc, ref, after=100)

    doc.save(OUT_DOCX)
    print("Wrote %s" % os.path.abspath(OUT_DOCX))
    if MISSING:
        print("\n*** %d FIGURE(S) NOT FOUND -- placeholders written in red:"
              % len(MISSING))
        for m in MISSING:
            print("      %s" % m)
    if REVIEW_NOTES:
        print("\n*** DATA NOTES:")
        for n in REVIEW_NOTES:
            print("      %s" % n)
    print("\nSection 13 is a placeholder for the downstream extension.")


main()
