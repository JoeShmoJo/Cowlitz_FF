#make_memo_combined.py
# -*- coding: utf-8 -*-
"""
Build the combined unregulated + regulated flow frequency memorandum.

This is a SEPARATE document from the two it draws on -- MEMO_CAS_Unreg_FF_
04Aug2026.docx and MEMO_CAS_Reg_Unreg_DRAFT.docx. Those two originals are left
untouched; this script never opens them for writing, only as a style/figure
source. The combined memo exists because the two originals repeat a lot of the
same ground (basin description, previous-studies note, data table entries,
references) and a reviewer flagged the duplication.

Style, page setup, and table formatting are taken from the adopted unregulated
memo, the same way make_memo_reg_unreg.py does it: the template is copied, its
body is emptied, and new content is written back with the same style names and
direct formatting the original uses, so the combined memo reads as part of the
same series.

Regulated-side ordinates and uncertainty terms are pulled from
regulated_frequency_inferred.csv, so this always reflects whatever
#Unreg_Reg_Curve.py last computed (currently the 95% band on the reviewer's
formula, aligned to EM 1110-2-1619 -- see Section 5.6 below). Unregulated-side
numbers are NOT pulled from a CSV -- there is no generator script for that
memo in this repo, the numbers were already reviewed once, and re-deriving
them here risks silent drift from what was approved. They are transcribed
directly from the adopted unregulated memo's tables.

Figures are the actual images embedded in the two source memos (extracted to
figures_combined/), NOT re-rendered -- except the final uncertainty figure,
which is pulled fresh from the diagnostics output because the embedded one in
the regulated DRAFT still shows the old 90% band.

Voice: tightened relative to both originals, more in the regulated memo than
the unregulated one (the user had already revised the unregulated memo's
wording once). The aim is a reviewer who knows hydrology already -- state what
was done and the adopted numbers, keep the reasoning that justifies a
methodological choice, drop the reasoning that just explains a term a
hydrologist already knows.

DRAFT 2. A COPY of make_memo_combined.py, writing DRAFT2.docx. The
original generator and MEMO_CAS_Combined_FlowFrequency_DRAFT.docx are
left untouched so the first draft stays available for comparison.
Section 8 is the only substantive change: it was [TO BE COMPLETED] and
is now written, with a new Appendix F carrying its ordinates.

!! THIS OVERWRITES OUT_DOCX !!  Once this draft is under review, hand edits in
Word are the source of record and this script becomes history, the same as
its regulated-only sibling.
"""

import os
os.chdir(os.path.dirname(os.path.abspath(__file__)))

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _appendix_data import APPENDIX_A_ROWS, APPENDIX_B_ROWS, APPENDIX_D_ROWS

TEMPLATE_DOCX = r"../../CAS_Unreg_FF/docs/MEMO_CAS_Unreg_FF_04Aug2026.docx"
OUT_DOCX = r"MEMO_CAS_Combined_FlowFrequency_DRAFT2.docx"
PRISM_CSV = r"../output/diagnostics/prism_basin_precip_ratio.csv"
# Section 8 / Appendix F, from #BelowConfluence_FlowFrequency.py
BELOW_CONF_CSV = r"../output/below_confluence_frequency.csv"

REG_FREQ_CSV = r"../output/regulated_frequency_inferred.csv"

FIG = {
    "basin": r"figures_combined/basin.png",
    "regression": r"figures_combined/regression_2day.png",
    "unreg_curves": r"figures_combined/unreg_adopted_curves.png",
    "critdur": r"figures_combined/critical_duration.png",
    "synth_events": r"figures_combined/synthetic_events.png",
    "scatter": r"figures_combined/transform_scatter.png",
    "final": r"figures_combined/final_uncertainty.png",
    "unreg_2009": r"figures_combined/unreg_vs_2009.png",
    "reg_2009": r"figures_combined/reg_vs_2009.png",
    "reg_2009_pct": r"figures_combined/reg_vs_2009_pctchange.png",
    "curve_peak": r"figures_combined/unreg_curve_peak.png",
    "curve_1day": r"figures_combined/unreg_curve_1day.png",
    "curve_3day": r"figures_combined/unreg_curve_3day.png",
    "curve_5day": r"figures_combined/unreg_curve_5day.png",
    "below_conf": r"figures_combined/below_confluence_frequency.png",
    "freq_gage": r"figures_combined/freq_castle_rock_gage.png",
    "freq_arkansas": r"figures_combined/freq_below_arkansas_creek.png",
    "freq_ostrander": r"figures_combined/freq_below_ostrander_creek.png",
    "freq_coweeman": r"figures_combined/freq_below_coweeman_river.png",
    "convergence": r"figures_combined/transform_convergence.png",
}
FIG_WIDTH_IN = 6.5

SZ_TABLE = 19
SZ_CAPTION = 19
SZ_H1, SZ_H2 = 26, 23
SZ_TITLE = 28
HDR_FILL = "D9E2EC"
BORDER_OUTER, BORDER_INNER = "7A869A", "B7C0CC"
TABLE_WIDTH_DXA = 9360

REVIEW_NOTES = []
MISSING = []

# ----------------------------------------------------------------------------
# style helpers -- identical to make_memo_reg_unreg.py, so the combined memo
# formats the same way its two sources do.


def para(doc, text="", size=None, bold=False, italic=False, align=None,
         style=None, before=None, after=160, line=276):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = Pt(before / 20.0)
    if after is not None:
        pf.space_after = Pt(after / 20.0)
    if line:
        pf.line_spacing = line / 240.0
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size / 2.0)
    return p


def h1(doc, text):
    return para(doc, text, size=SZ_H1, bold=True, style="Heading 1",
                before=300, after=140, line=None)


def h2(doc, text):
    return para(doc, text, size=SZ_H2, bold=True, style="Heading 2",
                before=300, after=140, line=None)


def caption(doc, text):
    return para(doc, text, size=SZ_CAPTION, italic=True,
                align=WD_ALIGN_PARAGRAPH.CENTER, before=80, after=240,
                line=None)


def figure(doc, key, cap, width_in=FIG_WIDTH_IN):
    path = FIG.get(key)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    if path and os.path.isfile(path):
        p.add_run().add_picture(path, width=Inches(width_in))
    else:
        run = p.add_run("[FIGURE NOT FOUND: %s]" % (path or key))
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2D)
        MISSING.append(path or key)
    caption(doc, cap)


def _borders(tbl):
    """Append w:tblBorders to tblPr in schema order.

    make_memo_reg_unreg.py's version of this helper just appends tblBorders
    at the end of tblPr, which lands it AFTER w:tblLayout (added by the
    autofit=False assignment made just before this runs in table()) --
    CT_TblPrBase requires tblBorders before tblLayout. Word opens the result
    anyway, but it fails strict OOXML validation and LibreOffice in this
    sandbox refuses to load it at all (confirmed: the existing
    MEMO_CAS_Reg_Unreg_DRAFT.docx has the same defect and the same
    soffice failure -- this is not new). Inserting immediately before
    tblLayout, when present, fixes it.
    """
    pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge, size, color in (("top", 4, BORDER_OUTER), ("left", 4, BORDER_OUTER),
                              ("bottom", 4, BORDER_OUTER), ("right", 4, BORDER_OUTER),
                              ("insideH", 2, BORDER_INNER), ("insideV", 2, BORDER_INNER)):
        el = OxmlElement("w:%s" % edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    layout = pr.find(qn("w:tblLayout"))
    if layout is not None:
        layout.addprevious(borders)
    else:
        pr.append(borders)


def _shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def table(doc, header, rows, widths=None, align_right=(), font_size=None):
    tbl = doc.add_table(rows=1, cols=len(header))
    tbl.style = doc.styles["Normal Table"]
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    _borders(tbl)

    if widths is None:
        widths = [TABLE_WIDTH_DXA // len(header)] * len(header)
    sz = font_size or SZ_TABLE

    def write(cell, text, bold, idx):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        if bold:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif idx in align_right:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(sz / 2.0)

    for i, text in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.width = Inches(widths[i] / 1440.0)
        _shade(cell, HDR_FILL)
        write(cell, text, True, i)
    for row in rows:
        cells = tbl.add_row().cells
        for i, text in enumerate(row):
            cells[i].width = Inches(widths[i] / 1440.0)
            write(cells[i], text, False, i)
    return tbl


def clear_body(doc):
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect:
            body.remove(child)
    return doc


def assert_no_comment_parts(path):
    import zipfile
    with zipfile.ZipFile(path) as archive:
        bad = [n for n in archive.namelist() if "comments" in n.lower()]
    if bad:
        raise SystemExit(
            "TEMPLATE_DOCX carries comment parts and cannot be used as a "
            "template:\n   %s\n   %s" % (path, ", ".join(bad)))


def fmt(value, digits=0):
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return "—"
    return format(round(float(value), digits), ",.%df" % digits)


def main():
    reg = pd.read_csv(REG_FREQ_CSV)
    conf = pd.read_csv(BELOW_CONF_CSV).sort_values(
        "AEP", ascending=False)
    prism = pd.read_csv(PRISM_CSV)
    if reg["freq_term_mode"].nunique() != 1:
        REVIEW_NOTES.append("regulated_frequency_inferred.csv mixes "
                            "freq_term_mode values -- Section 5.6 text "
                            "assumes one mode throughout.")
    freq_mode = reg["freq_term_mode"].iloc[0]

    def at(aep, col):
        i = (reg["AEP"] - aep).abs().idxmin()
        v = reg.loc[i, col]
        return None if pd.isna(v) else float(v)

    assert_no_comment_parts(TEMPLATE_DOCX)
    doc = clear_body(Document(TEMPLATE_DOCX))

    # ---------------------------------------------------------------- header
    para(doc, "MEMORANDUM FOR RECORD", size=SZ_TITLE, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=60, line=None)
    para(doc, "U.S. Army Corps of Engineers, Portland District", size=21,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=240, line=None)
    table(doc, ["Item", "Description"], [
        ["Subject", "Unregulated and Regulated Flow Frequency, Cowlitz "
                    "River at Castle Rock, Washington (USGS 14243000)"],
        ["From", "Hydrology and Hydraulics Branch, Portland District"],
        ["Date", "August 2026"],
        ["Status", "DRAFT — combines the separate unregulated and regulated "
                   "memoranda into one document; Section 8 (Columbia River "
                   "confluence extension) pending"],
    ], widths=[2400, 6960])
    para(doc, "", after=120)

    # ------------------------------------------------------------- 1 Purpose
    h1(doc, "1. Purpose")
    para(doc, "This memorandum documents the development of unregulated and "
              "regulated flow frequency curves for the Cowlitz River at "
              "Castle Rock, Washington. Unregulated curves are developed for "
              "the instantaneous peak and the 1-, 3-, and 5-day durations "
              "using water years 1927 through 2026 (Section 4); the "
              "regulated curve for the instantaneous peak is then obtained "
              "by transforming the adopted unregulated peak curve through a "
              "relationship between unregulated and regulated flow derived "
              "from reservoir simulation, rather than by fitting a "
              "distribution to the observed regulated record (Section 5). "
              "The regulated curve, with its uncertainty, supports levee "
              "fragility assessment and floodplain mapping in the lower "
              "Cowlitz valley.")
    para(doc, "Section 8 carries the regulated curve downstream from the "
              "gage to three additional locations — below Arkansas Creek, "
              "below Ostrander Creek, and below the Coweeman River — by "
              "adding the contribution of the intervening local drainage. "
              "Those four locations are the deliverable; no analysis below "
              "the Coweeman confluence is included.")
    para(doc, "Neither curve can be produced by fitting the observed record "
              "directly. Regulated peaks are not a homogeneous sample — "
              "operating rules and starting-pool practice have changed since "
              "Mossyrock began impounding in 1968 — and they do not follow "
              "an analytical distribution: a reservoir holds small events "
              "almost entirely, loses ground through the middle of its "
              "range, and passes the largest events through once storage is "
              "exhausted, a shape no log-Pearson Type III curve reproduces. "
              "The standard approach, used here, is to establish the "
              "frequency of the natural inflow and carry it through the "
              "projects.")

    # ------------------------------------------------- 2 Basin and operation
    h1(doc, "2. Basin Description and Project Operation")
    para(doc, "The Cowlitz River drains the southern Washington Cascades and "
              "enters the Columbia River near Longview. The Castle Rock gage "
              "(USGS 14243000) sits in the lower valley, downstream of the "
              "Toutle River confluence and of all the basin's storage "
              "projects, with a peak-flow record spanning water years 1927 "
              "through 2026. The total basin area above Castle Rock is "
              "2,238 square miles, of which 1,170 square miles of headwaters "
              "are regulated by Mossyrock Dam.")
    para(doc, "Three non-Federal hydroelectric projects lie on the mainstem, "
              "operated by Tacoma Power as the Cowlitz Hydroelectric "
              "Project. Mossyrock Dam, completed in 1968 and impounding "
              "Riffe Lake, provides the only usable flood storage in the "
              "basin and began actively regulating the river in December "
              "1968; years before then are unregulated as observed. "
              "Mayfield Dam, completed in 1963, is a re-regulating dam that "
              "smooths power peaking releases from Mossyrock and provides "
              "no meaningful flood regulation of its own. Cowlitz Falls Dam, "
              "above Mossyrock, and the Sediment Retention Structure on the "
              "Toutle are both run-of-river.")
    para(doc, "Two features of the basin govern how much the projects can "
              "do, and both are visible in the results that follow. The "
              "regulated fraction is a little over half the drainage area, "
              "so roughly half the flow arriving at Castle Rock in a large "
              "event is uncontrolled. The Toutle River, the largest "
              "tributary, enters below both dams and is entirely "
              "unregulated; the Tilton River enters Mayfield Lake, which "
              "passes inflow and provides no flood regulation. At the "
              "largest events the local and Toutle contribution is a median "
              "48 percent of unregulated flow at Castle Rock, which sets a "
              "floor on the regulated peak that no operation can reach "
              "below.")
    figure(doc, "basin", "Figure 2-1. Cowlitz River basin, showing the "
                         "Castle Rock gage, the Mossyrock and Mayfield "
                         "projects, and the Toutle and Tilton tributaries.")

    h2(doc, "2.1 Previous Studies")
    para(doc, "The previously adopted unregulated and regulated frequency "
              "curves at Castle Rock both come from the 2009 Cowlitz "
              "Hydrology Restudy, which predates Bulletin 17C, does not "
              "include the most recent 17 years of record, and derived its "
              "regulated curve by a different route than the one used here. "
              "Section 6 compares both curves against the 2009 study.")

    # ------------------------------------------------------------ 3 Data used
    h1(doc, "3. Data Used")
    para(doc, "Table 3-1 lists the datasets and models used and the purpose "
              "each served. The observed hydrologic records are shared by "
              "both analyses; the reservoir simulation model and its "
              "outputs are specific to the regulated study.")
    table(doc, ["Dataset or model", "Period used", "Purpose"], [
        ["Mossyrock pool elevation, hourly", "Nov 1974 – May 2026",
         "Reservoir storage change at hourly resolution; basis of the "
         "hourly unregulated peak estimates (5.2)"],
        ["Mossyrock storage telemetry, hourly", "Regulated era",
         "Data quality screening"],
        ["Mossyrock pool elevation, daily", "Oct 1973 – present",
         "Storage change for the regression predictors and the daily mass "
         "balance (5.3, 5.4); observed starting pool for the Obs_RC "
         "simulation (5.1)"],
        ["Mossyrock elevation-storage rating table", "2014 Water Control "
         "Manual", "Relates pool elevation to reservoir storage"],
        ["Castle Rock streamflow, hourly (USGS 14243000)", "Regulated era",
         "Regulated flow at the study location; source of the regulated "
         "peak"],
        ["Castle Rock streamflow, daily (USGS 14243000)", "WY1927 – present",
         "Pre-regulation durations; regulated daily flow used for the "
         "3-, 5-, and select 1-day duration maxima"],
        ["Castle Rock annual peak flows (USGS 14243000)", "WY1927 – present",
         "Pre-regulation peaks; regulated peak for years the hourly record "
         "does not cover; basis of the adjusted regulated peak record "
         "(5.1)"],
        ["Mayfield outflow, hourly (USGS 14238000)", "Regulated era",
         "Upstream boundary for routing the Mossyrock regulation signal to "
         "Castle Rock"],
        ["HEC-ResSim model, Mossyrock and Mayfield", "Current (2014 WCM) "
         "rules", "Routes inflow hydrographs through the projects for the "
         "adjustment (5.1) and the synthetic ensembles (5.3)"],
        ["Reservoir inflow and Castle Rock local flow", "WY 1929 – 2026",
         "Simulation inflows, back-calculated and volume-corrected"],
        ["ResSim period-of-record simulation, rule-curve start (WCM_RC)",
         "98 water years", "Regulated peak under current rules from a "
         "rule-curve starting pool"],
        ["ResSim period-of-record simulation, observed start (Obs_RC)",
         "53 water years", "Regulated peak under current rules from the "
         "observed starting pool"],
        ["ResSim synthetic ensemble simulation", "48 members",
         "Regulated response above the observed range (5.3)"],
    ], widths=[3600, 1700, 4060])
    caption(doc, "Table 3-1. Datasets and models used, and their purpose.")

    # ============================================================
    # 4. UNREGULATED FLOW FREQUENCY
    # ============================================================
    h1(doc, "4. Unregulated Flow Frequency")
    para(doc, "Unregulated flow is the flow that would have occurred at "
              "Castle Rock without Mossyrock. Because the reservoir changes "
              "downstream flow only by storing and releasing water, the "
              "unregulated hydrograph is reconstructed by adding back the "
              "holdout — reservoir storage change expressed as a flow rate "
              "— to the observed record. A positive holdout is the "
              "reservoir filling; a negative holdout is releasing more "
              "than inflow.")
    para(doc, "For water years 1927–1968 the gage record is unregulated by "
              "definition and used as observed (Section 4.4). For the "
              "regulated period the reconstruction runs at whatever "
              "resolution the reservoir record supports: an hourly holdout "
              "where the hourly elevation record is clean enough (Section "
              "4.1), otherwise a peak estimated by regression from the "
              "daily reservoir record (Section 4.2). The 3- and 5-day "
              "durations, needing only daily resolution, come from a daily "
              "mass balance in all years (Section 4.3).")

    h2(doc, "4.1 Hourly Unregulated Record")
    para(doc, "The cleaned hourly pool elevation was converted to storage "
              "using the 2014 Water Control Manual rating curve and "
              "differenced to obtain an hourly holdout, smoothed with a "
              "3-hour centered moving average and restricted to the "
              "October–March flood season. Days with fewer than four valid "
              "storage readings were omitted, since they cannot be relied "
              "on to capture the annual peak.")
    para(doc, "The reservoir signal was translated to Castle Rock using the "
              "SSARR routing parameters in Table 4-1, from the calibrated "
              "Cowlitz CWMS model. Mayfield outflow plus the holdout, and "
              "Mayfield outflow alone, were each routed to Castle Rock; the "
              "difference between the two routed hydrographs is the routed "
              "reservoir effect, added to observed Castle Rock flow to give "
              "the hourly unregulated estimate. The Toutle River reach was "
              "excluded because only the mainstem reservoir signal is being "
              "routed — Toutle inflow is already in the observed Castle "
              "Rock record. Mayfield-to-Castle Rock travel time is short "
              "enough that routing barely matters there, but all three "
              "reaches carry parameters for completeness.")
    table(doc, ["Reach", "KTS", "n", "Phases"], [
        ["Mayfield outflow to Cowlitz River above Toutle River", "5", "0.1", "5"],
        ["Cowlitz River above Toutle River to Cowlitz River below Toutle River",
         "1", "0.2", "1"],
        ["Cowlitz River below Toutle River to Castle Rock", "1", "0.2", "5"],
    ], widths=[6360, 1000, 1000, 1000], align_right=(1, 2, 3))
    caption(doc, "Table 4-1. SSARR routing parameters, from the calibrated "
                "lower Cowlitz routing model.")
    para(doc, "Twenty-three water years produced holdout-based peaks — "
              "1992, 1997, 1999, 2003, 2004, and 2008 through 2026 — with "
              "earlier hourly pool elevation too unstable for use.")

    h2(doc, "4.2 Peak Estimation by Storage-Change Regression")
    para(doc, "For years without a usable hourly record — roughly a third "
              "of the regulated era — the unregulated peak is instead "
              "estimated by regressing the regulated-minus-unregulated peak "
              "difference on the largest 1- to 4-day increase in Mossyrock "
              "storage near the event (within 7 days of the regulated "
              "peak), then subtracting the predicted difference from the "
              "observed regulated peak. This is preferred over regressing "
              "the unregulated peak directly: the predictor and the "
              "regulated peak are both available for every regulated year, "
              "and because the adjustment is a fraction of the flow being "
              "reconstructed, a given percentage error on it translates to "
              "a much smaller error on the resulting unregulated peak.")
    table(doc, ["Storage-change window", "Slope", "Intercept (cfs)",
               "R²", "Standard error (cfs)"], [
        ["1-day", "-0.673", "-7,751", "0.759", "7,565"],
        ["2-day (adopted)", "-0.869", "-3,836", "0.871", "5,528"],
        ["3-day", "-0.983", "-4,537", "0.853", "5,896"],
        ["4-day", "-1.140", "-4,770", "0.836", "6,239"],
    ], widths=[3000, 1600, 1900, 1400, 1460], align_right=(1, 2, 3, 4))
    caption(doc, "Table 4-2. Candidate regressions of the regulated-minus-"
                "unregulated peak difference on Mossyrock storage change, "
                "fit to 17 water years.")
    para(doc, "The two-day window was adopted — highest R², lowest "
              "standard error: peak difference (regulated minus "
              "unregulated) = −0.869 × two-day storage change "
              "− 3,836 cfs, R² = 0.871, standard error ≈ 5,530 "
              "cfs, n = 17. Pairs whose regulated and unregulated peaks "
              "fell more than 72 hours apart were excluded as different "
              "storms. Thirty water years were filled by regression; for "
              "1974–1991, where the hourly record does not reach, the "
              "regulated peak came from the USGS annual peak record and "
              "the peak date centered the storage-change window.")
    figure(doc, "regression", "Figure 4-1. Adopted two-day storage-change "
                              "regression. Points are labeled by water "
                              "year.")

    h2(doc, "4.3 Durations from the Daily Mass Balance")
    para(doc, "Daily unregulated flow is observed daily flow plus the daily "
              "change in Mossyrock storage as a flow rate; no routing is "
              "applied since Mayfield-to-Castle Rock travel time is under a "
              "day, and Mayfield re-regulation storage is neglected for the "
              "same reason. Rolling 1-, 3-, and 5-day maxima were computed "
              "on a complete daily grid so an N-day value always reflects N "
              "consecutive recorded days. The Castle Rock daily record runs "
              "winter-season-only through much of the 1980s–1990s, which is "
              "taken to capture the annual maximum since it spans the flood "
              "season. The one-day mass balance value also supplies the "
              "1-day duration in the thirty regulated years without a "
              "usable hourly record.")

    h2(doc, "4.4 Pre-Regulation Period")
    para(doc, "Before December 1968 the gage record is unregulated as "
              "observed. Peaks for water years 1927–1968 are the USGS "
              "observed instantaneous annual peaks (42 water years), and "
              "the 1-, 3-, and 5-day values are rolling maxima from the "
              "USGS daily record. Water year 1927 carries a peak only; the "
              "daily record begins later.")

    h2(doc, "4.5 Record Extension and Historical Floods")
    para(doc, "Correlation-based extension from a nearby gage was "
              "considered but no suitable donor was found — nearby "
              "long-record gages either do not predate Castle Rock by much, "
              "are themselves regulated over the period that would be "
              "used, or do not correlate strongly enough to justify the "
              "added estimation error; with roughly a century of systematic "
              "record already in hand, the potential gain was small. A "
              "search of USGS qualification codes, the 2009 Restudy, and "
              "published USGS/NWS accounts of major southwest Washington "
              "floods turned up no pre-1927 event at Castle Rock that could "
              "be quantified or reliably bounded, so no historical events "
              "or Bulletin 17C perception thresholds were used.")

    h2(doc, "4.6 Assembled Record")
    para(doc, "For water years 1927–1968 the record is the observed gage "
              "record; for 1974–2026 the peak is the hourly holdout where "
              "available and the regression estimate otherwise, the 1-day "
              "value follows the same split, and the 3- and 5-day values "
              "always come from the daily mass balance (Table 4-3, "
              "Appendix A). Coverage runs from WY1927 through WY2026 except "
              "for WY1969–1973, which have no daily Mossyrock elevation "
              "record — it begins October 1973 — and so no basis for a "
              "holdout or a regression predictor. Every included water year "
              "has a peak; WY1927 has no durations since the daily record "
              "starts later.")
    table(doc, ["Duration", "Values", "Sources"], [
        ["Peak", "95", "42 pre-regulation USGS peaks; 30 storage-change "
                       "regression; 23 hourly holdout"],
        ["1-Day", "94", "41 pre-regulation USGS daily; 30 daily mass "
                        "balance; 23 hourly holdout"],
        ["3-Day", "94", "41 pre-regulation USGS daily; 53 daily mass "
                        "balance"],
        ["5-Day", "94", "41 pre-regulation USGS daily; 53 daily mass "
                        "balance"],
    ], widths=[1400, 1000, 6960], align_right=(1,))
    caption(doc, "Table 4-3. Assembled record, values and sources by "
                "duration.")

    h2(doc, "4.7 Consistency Review")
    para(doc, "Each duration is the independent annual maximum of its own "
              "averaging window, so nothing in the construction forces a "
              "shorter duration to exceed a longer one — yet a "
              "shorter-duration flow lower than a longer-duration flow for "
              "a single event is physically impossible. Fifteen such "
              "violations appeared, fourteen from a source mismatch "
              "(typically a regression-estimated peak below the same "
              "year's mass-balance 1-day value) and one, WY2014, from a "
              "double-peaked March event whose 5-day window spans both "
              "peaks while no 3-day window can. The record was made "
              "monotonic by anchoring each water year at its 5-day value "
              "and raising each shorter duration wherever a longer one "
              "exceeds it, cascading upward; pre-adjustment values are "
              "retained alongside the adopted ones. Ten values across nine "
              "water years were raised — eight of them regression-"
              "estimated peaks lifted to their 1-day value — with the "
              "largest changes, 64, 19, and 19 percent, in WY2001, 1985, "
              "and 2006, all among the smallest years in the record and so "
              "of little influence on the fitted upper tail.")

    h2(doc, "4.8 Flow Frequency Analysis")
    para(doc, "The assembled record was analyzed in HEC-SSP with the "
              "Bulletin 17C expected moments algorithm, fitting a "
              "log-Pearson Type III distribution to each duration and "
              "weighting station skew against a published regional value: "
              "−0.07 for the peak (Mastin et al., 2016, a constant "
              "regional skew for a Pacific Northwest region including the "
              "Cowlitz basin), and 0.09 / 0.07 / 0.00 for the 1-, 3-, and "
              "5-day durations (Lind et al., 2020, Bayesian regional skew "
              "models for Columbia River basin durations, driven by basin "
              "mean annual precipitation — the 5-day value is interpolated "
              "between that study's bracketing 3- and 7-day models at a "
              "basin-average 77.5 inches).")
    para(doc, "Water years 1969–1973 have observed regulated peaks but no "
              "unregulated basis, so rather than omit them (shortening the "
              "record) or assign point estimates (asserting unsupported "
              "values), they were entered as censored: the peak analysis "
              "reports 95 systematic events within a 100-year period, with "
              "five flagged missing. No low outliers were identified by "
              "the multiple Grubbs-Beck procedure — the unregulated series "
              "is constructed rather than gaged, so its small years carry "
              "the same provenance as its large ones and there is no "
              "physical basis for treating the lower tail as "
              "unrepresentative.")
    para(doc, "Station skews are slightly negative for every duration "
              "(−0.158 to −0.054); weighting against the regional "
              "values gives adopted skews of −0.135, −0.052, "
              "−0.009, and −0.031 for the peak, 1-, 3-, and 5-day "
              "durations, with equivalent record lengths of 100–112 years "
              "after censoring. Table 4-4 gives the fitted parameters and "
              "event counts; Figure 4-2 shows the four adopted curves. "
              "Individual curves with confidence limits and plotted "
              "observations appear in Appendix C, and the full ordinates "
              "in Appendix B.")
    table(doc, ["Parameter", "Peak", "1-Day", "3-Day", "5-Day"], [
        ["Mean of logs", "4.778", "4.743", "4.660", "4.593"],
        ["Standard deviation of logs", "0.202", "0.203", "0.192", "0.179"],
        ["Station skew", "-0.158", "-0.152", "-0.076", "-0.051"],
        ["Regional skew", "-0.07", "0.09", "0.07", "0.00"],
        ["Weighted skew", "-0.135", "-0.052", "-0.007", "-0.029"],
        ["Adopted skew", "-0.135", "-0.052", "-0.009", "-0.031"],
        ["Mean square error, at-site skew", "0.064", "0.064", "0.060", "0.059"],
        ["Grubbs-Beck critical value", "0", "0", "0", "0"],
        ["Equivalent record length (years)", "103.7", "100.4", "106.3", "112.0"],
        ["Systematic events", "95", "94", "94", "94"],
        ["Historical events", "0", "0", "0", "0"],
        ["High outliers", "0", "0", "0", "0"],
        ["Low outliers", "0", "0", "0", "0"],
        ["Zero events", "0", "0", "0", "0"],
        ["Missing events", "5", "5", "5", "5"],
        ["Historical period (years)", "100", "99", "99", "99"],
    ], widths=[3760, 1400, 1400, 1400, 1400], align_right=(1, 2, 3, 4))
    caption(doc, "Table 4-4. Distribution parameters and event counts, "
                "adopted HEC-SSP analyses.")
    figure(doc, "unreg_curves", "Figure 4-2. Adopted unregulated frequency "
                                "curves at Castle Rock, water years 1927 "
                                "through 2026, all four durations.")

    # ============================================================
    # 5. REGULATED FLOW FREQUENCY
    # ============================================================
    h1(doc, "5. Regulated Flow Frequency")
    para(doc, "Regulated peaks cannot be fit directly, for two reasons: the "
              "observed record is not a homogeneous sample — operating "
              "rules and starting-pool practice have varied since Mossyrock "
              "began impounding in 1968 — and regulated peaks do not follow "
              "an analytical distribution, since operating rules put hard "
              "breaks in the inflow–outflow relationship that no "
              "log-Pearson Type III curve reproduces. The regulated curve "
              "is instead built by adjusting the observed peak record for "
              "the starting-pool effect (5.1), pairing adjusted peaks with "
              "the unregulated peak of the same event (5.2), augmenting "
              "the upper end with routed synthetic floods (5.3), and "
              "drawing a transform through the combined set that is "
              "applied to the adopted unregulated curve (5.4–5.5).")

    h2(doc, "5.1 Adjusted Regulated Peak Record")
    para(doc, "The observed Mossyrock record shows the reservoir frequently "
              "drafted below the flood-pool minimum, so a peak from a year "
              "that started low understates the attenuation the project "
              "would provide from its full flood-control authority. Two "
              "period-of-record ResSim simulations, using the same "
              "observed hydrology and current release rules and differing "
              "only in starting pool — one from the WCM rule curve, one "
              "from the observed elevation at event onset — isolate that "
              "effect. Where the rule-curve start yields the higher "
              "regulated peak, the difference is added to the observed "
              "peak; the adjustment is one-sided, since a negative "
              "difference means the historical operation was already at "
              "least as conservative as the rule curve requires.")
    para(doc, "adjusted peak = observed peak + (simulated RC-start peak "
              "− simulated observed-start peak), where the RC-start "
              "peak is higher", italic=True, before=40, after=120, line=None)
    para(doc, "The period of record is 1974 onward, set by the daily "
              "Mossyrock elevation record (beginning October 1973). Peaks "
              "are matched by event, not water year, so that both "
              "simulated peaks, the observed peak, and the unregulated peak "
              "all reflect the same storm; the annual max is used where "
              "that holds.")
    para(doc, "Of 51 shared water years (1974–2024), 38 received an "
              "adjustment (median 11,334 cfs, range 263–24,013). Screening "
              "removed 10: 7 for a timing mismatch between the observed and "
              "simulated peaks, 2 where the regulated peak exceeded the "
              "unregulated peak (physically only possible if release "
              "exceeds inflow during a large flood), and 1 for both. "
              "WY1980 is excluded outright — its peak is the Mount St. "
              "Helens lahar, not a hydrologic flood.")

    h2(doc, "5.2 Critical Duration")
    para(doc, "Mapping a regulated flow onto the unregulated curve requires "
              "choosing which unregulated duration corresponds to the "
              "regulated peak (Table 5-1, Figure 5-1). Both the "
              "instantaneous peak and the 1-day duration are strong "
              "predictors; the median ratio explains why the shorter "
              "durations correlate better — at the peak the projects hold "
              "flow to about 78 percent of unregulated, while at 5 days "
              "the ratio exceeds one, because water held out of the peak "
              "is released afterward and reappears in the longer window. "
              "The projects move flood volume in time, not out of it. The "
              "peak unregulated flow is adopted for the transform.")
    table(doc, ["Duration", "n", "R²", "Log R²", "Log slope",
               "Median reg / unreg"], [
        ["Peak (1-hr)", "44", "0.870", "0.847", "0.602", "0.78"],
        ["1-Day", "44", "0.853", "0.830", "0.605", "0.87"],
        ["3-Day", "44", "0.792", "0.783", "0.633", "1.15"],
        ["5-Day", "44", "0.733", "0.734", "0.663", "1.45"],
    ], widths=[2960, 1000, 1200, 1400, 1400, 2400], align_right=(1, 2, 3, 4, 5))
    caption(doc, "Table 5-1. Adjusted regulated against unregulated flow by "
                "duration, water years 1974 through 2024.")
    figure(doc, "critdur", "Figure 5-1. Adjusted regulated against "
                          "unregulated flow, by duration.")

    h2(doc, "5.3 Synthetic Flood Ensembles")
    para(doc, "Only one water year in the record exceeds the unregulated "
              "100-year peak, so the part of the unregulated-regulated "
              "relationship that sets the 100-, 250-, and 500-year "
              "regulated flows is essentially unconstrained by observed "
              "events. Synthetic floods fill that range: magnitude is "
              "varied over four targets from the unregulated curve (the "
              "100-, 250-, and 500-year peaks, plus a fourth 20 percent "
              "above the 500-year so the fit is not extrapolating exactly "
              "where the answer is needed), and shape is varied by scaling "
              "twelve observed storms spanning the full range of observed "
              "attenuation behavior — shape is varied rather than assumed "
              "because attenuation among large events is not predictable "
              "from magnitude, antecedent flow, or timing alone.")
    para(doc, "Each source storm is scaled to hit its target unregulated "
              "peak and 5-day volume (within 5 percent), with Mossyrock "
              "inflow and Castle Rock local scaled by the same factor to "
              "preserve the observed coincidence between the controlled "
              "and uncontrolled fractions. The 48 routed members span "
              "regulated peaks of 80,884–266,744 cfs and attenuation "
              "ratios of 0.46–1.00; the spread at a single magnitude is a "
              "shape effect, not scatter — Nov2006 is absorbed most "
              "strongly, Dec1933 exhausts storage and passes through near "
              "1:1 — and is the physical basis for the uncertainty band in "
              "Section 5.6.")
    figure(doc, "synth_events", "Figure 5-2. Source storms and the scaled "
                                "family built from each. The shaded band "
                                "is the window over which the 5-day volume "
                                "is matched.")

    h2(doc, "5.4 The Unregulated-Regulated Transform")
    para(doc, "The adjusted historical pairs (observed range) and the "
              "routed synthetic members (range above it) together give the "
              "(unregulated, regulated) pairs the transform is drawn "
              "through: a locally weighted regression in log-log space, "
              "tricube-weighted over the nearest 65 percent of the sample, "
              "constrained to increase monotonically and clipped at the "
              "1:1 line, since a regulated peak cannot exceed its own "
              "unregulated peak in a large event.")
    para(doc, "The WCM_RC simulated pairs are shown for reference but not "
              "fitted — they are simulated on both axes with no observed "
              "anchor, since the historical-record adjustment (5.1) is "
              "built from the difference between the two simulations, and "
              "where no Obs_RC data exists there is no correction to "
              "apply. They speak to the shape of the relationship, not its "
              "position.")
    figure(doc, "scatter", "Figure 5-3. Unregulated against regulated peak "
                          "at Castle Rock, log-log, with the adopted "
                          "transform and the single power law for "
                          "reference.")

    h2(doc, "5.5 Regulated Frequency Curve")
    para(doc, "Each ordinate of the adopted unregulated curve is read "
              "through the transform, inheriting its AEP: the regulated "
              "flow at a given AEP is what the projects produce from the "
              "unregulated flow of that AEP (Table 5-2). Mossyrock reduces "
              "the peak about 37 percent near the 1–2 percent AEP; the "
              "reduction is smaller at more frequent flows, where events "
              "are already largely contained, and at rarer flows, where "
              "storage is exhausted and the project must release more to "
              "avoid overtopping — down to about 30 percent at the 0.2 "
              "percent AEP.")
    sel_events = [("2-year", 0.500), ("10-year", 0.100), ("50-year", 0.020),
                  ("100-year", 0.010), ("200-year", 0.005), ("500-year", 0.002)]
    rows = []
    for label, aep in sel_events:
        rows.append([
            label, "%.3f" % aep,
            fmt(at(aep, "unreg_expected_cfs")),
            fmt(at(aep, "reg_inferred_cfs")),
            "%.0f%%" % at(aep, "reduction_pct"),
            "%s – %s" % (fmt(at(aep, "reg_lower_95pct_cfs")),
                        fmt(at(aep, "reg_upper_95pct_cfs"))),
        ])
    table(doc, ["Event", "AEP", "Unregulated (cfs)", "Regulated (cfs)",
               "Reduction", "Regulated, 95% band (cfs)"], rows,
        widths=[1360, 1000, 1800, 1800, 1200, 2200], align_right=(1, 2, 3, 4, 5))
    caption(doc, "Table 5-2. Adopted regulated frequency curve at Castle "
                "Rock, selected ordinates.")
    figure(doc, "final", "Figure 5-4. Adopted regulated and unregulated "
                        "frequency curves at Castle Rock, with the 95 "
                        "percent uncertainty band.")

    h2(doc, "5.6 Uncertainty")
    para(doc, "The regulated flow at a given AEP carries uncertainty from "
              "two independent sources — the unregulated frequency "
              "estimate itself, and the scatter of the transform about its "
              "fitted line — combined following EM 1110-2-1619 (29 Sep "
              "2025) Section 4-4.b(3): standard deviations from independent "
              "sources are combined by root-sum-of-squares (Eq. 4-6), and "
              "each source is treated as Normal on the upper and lower "
              "side of its best estimate separately (Section 4-6a), which "
              "preserves the real asymmetry in both the HEC-SSP confidence "
              "limits and the transform residuals rather than assuming one "
              "symmetric spread.")
    para(doc, "The adopted band is the reviewer's formula, at the 95 "
              "percent two-sided level (z = 1.960):", after=60)
    para(doc, "Upper = RegBest + √[(Unreg97.5 − Unreg_best)² + "
              "(Transform97.5 − Transform_best)²]", italic=True,
        before=0, after=20, line=None)
    para(doc, "Lower = RegBest − √[(Unreg_best − Unreg2.5)² + "
              "(Transform_best − Transform2.5)²]", italic=True,
        before=0, after=140, line=None)
    para(doc, "The frequency term is the HEC-SSP confidence-limit flow "
              "minus the best estimate, in unregulated cfs, taken directly "
              "rather than first pushed through the fitted transform curve "
              "— matching the reviewer's formula exactly. The two "
              "approaches differ by about 1 percent at the 1 percent AEP, "
              "growing to roughly 7 percent by the 0.1 percent AEP as the "
              "transform bends toward pass-through; close enough, and the "
              "reviewer's formula specific enough, that matching it "
              "directly was adopted (Appendix E carries both). The "
              "transform term is the LOESS fit's own upper/lower "
              "confidence spread, evaluated at the unregulated best "
              "estimate and already in regulated cfs.")
    para(doc, "Each side keeps its own sigma throughout (FREQ_TERM_MODE = "
              "%r in the analysis script). The SSP confidence limits are "
              "asymmetric — wider above the best estimate than below, "
              "more so at the rarer AEPs — and so, now, is the LOESS "
              "residual scatter: it is split into an upper- and "
              "lower-half sigma from the residuals on each side of the "
              "fit rather than pooled into one number. The transform "
              "scatter also varies along the curve and by side rather "
              "than being pooled — roughly 0.03 to 0.10 dex depending on "
              "magnitude and side — because whether a large flood "
              "exhausts storage depends on its shape while a small one is "
              "simply held." % freq_mode)
    para(doc, "The upper bound is held at or below the unregulated upper "
              "bound above 60,000 cfs, since a regulated flood cannot "
              "exceed the unregulated flood it was routed from once Castle "
              "Rock is above its 50,000–70,000 cfs regulation goal (below "
              "that threshold the project may release more than inflow to "
              "regain flood storage). The resulting band reaches about "
              "2.2× more lopsided on one side than the other at the "
              "rarest AEP, past the 2.0× threshold at which a two-piece "
              "lognormal should be checked against a Monte Carlo — that "
              "check has not yet been run. Appendix E gives every term by "
              "AEP.")

    # ============================================================
    # 6. COMPARISON WITH THE 2009 RESTUDY
    # ============================================================
    h1(doc, "6. Comparison with the 2009 Restudy")

    h2(doc, "6.1 Unregulated Curves")
    para(doc, "Figure 6-1 places the adopted unregulated curves against the "
              "2009 Cowlitz Hydrology Restudy; both are computed HEC-SSP "
              "curves. The two agree closely through the range supported "
              "by data — within about 2 percent at the 10 percent AEP, 3 "
              "percent at the median — and diverge in the tail: 7.8 "
              "percent lower at the 1 percent AEP (1.8–3.3 percent lower "
              "for the durations), 13.7 percent lower at the 0.2 percent "
              "AEP. The cause is skew — the 2009 study used a regional "
              "skew of 0 and adopted positive skews (0.140–0.193) that "
              "bend its curves upward in the tail, where the skews adopted "
              "here are near zero or negative and flatten it. The "
              "divergence is a difference in fitted tail behavior, not in "
              "the record through the observed range.")
    figure(doc, "unreg_2009", "Figure 6-1. Adopted 2026 unregulated curves "
                             "(solid) against the 2009 Cowlitz Hydrology "
                             "Restudy (dashed).")

    h2(doc, "6.2 Regulated Curve")
    para(doc, "Figure 6-2 places the adopted regulated curve against the "
              "2009 Restudy. The two agree within about 9 percent through "
              "the range supported by observed events — 13 percent above "
              "the 2009 value at the 1 percent AEP, 3 percent above at the "
              "0.2 percent AEP — then diverge, 21 percent below by the 0.1 "
              "percent AEP. The 2009 curve turns sharply upward above the "
              "0.2 percent AEP (156,000 to 390,000 cfs between the 0.2 and "
              "0.01 percent events), driven by synthetics built almost "
              "exclusively from the 1933 event, which fills and spills "
              "before its peak inflow; the current study simulates a "
              "broader set of plausible inflow shapes at similar volumes, "
              "most of which prove regulable.")
    figure(doc, "reg_2009", "Figure 6-2. Adopted 2026 regulated curve "
                           "against the 2009 Cowlitz Hydrology Restudy.")
    figure(doc, "reg_2009_pct", "Figure 6-3. Percent change between the "
                               "2026 and 2009 regulated curves.")

    # ============================================================
    # 7. LIMITATIONS
    # ============================================================
    h1(doc, "7. Limitations")
    para(doc, "The transform is supported by data over the unregulated "
              "range of roughly 23,000–280,000 cfs. Its upper part rests "
              "entirely on synthetic events built by scaling observed "
              "storms, and hydrograph shape is not independent of size — "
              "a large storm is large because it was widespread and "
              "sustained, so scaling a moderate storm to a rare magnitude "
              "routes a shape that storm never produced. Six of the "
              "twelve source storms are scaled past a factor of two to "
              "reach the largest target and should be read as the weaker "
              "evidence in the set.")
    para(doc, "The adjusted regulated record carries its own caveats: "
              "thirteen adjustments land within a narrow band near 14,000 "
              "cfs, more likely a fixed release constraint binding in one "
              "simulation and not the other than a pure starting-pool "
              "effect; seven adjustments exceed 40 percent of the observed "
              "peak, all moderate-peak years where a large absolute "
              "adjustment is a large fraction, influencing the lower end "
              "of the transform more than the upper.")
    para(doc, "The local contributions of Section 8 carry two further "
              "limitations. The drainage-area ratio is supported by "
              "observation only to about 212,000 cfs unregulated, the "
              "February 1996 event; beyond that it is an extrapolation, "
              "though a physically bounded one, since a basin cannot "
              "indefinitely exceed its area share once both it and the "
              "mainstem are saturated. And the 0.80 timing factor rests on "
              "events smaller than the design range: the seven largest "
              "events in the record give a substantially lower figure, and "
              "while that figure is not adopted for the reasons given in "
              "Section 8.3, it means the adopted value is supported by "
              "moderate events and applied to rare ones. Both limitations "
              "act in the conservative direction, and the sensitivity in "
              "Section 8.4 bounds their combined effect well inside the "
              "confidence band.")

    # ============================================================
    # 8. EXTENSION TO THE COLUMBIA RIVER CONFLUENCE
    # ============================================================
    h1(doc, "8. Extension Downstream to the Coweeman Confluence")
    para(doc, "The regulated frequency curve developed above applies at the "
              "Castle Rock gage, which drains 2,229 square miles. Between "
              "that gage and the Coweeman River confluence the Cowlitz "
              "picks up 247 square miles of additional drainage — Arkansas "
              "Creek, Ostrander Creek, the Coweeman River, and the ungaged "
              "local area along both banks. This section develops regulated "
              "peak frequency curves at each of those points.")

    h2(doc, "8.1 Method")
    para(doc, "At a site draining DA square miles, the regulated peak of "
              "annual exceedance probability p is taken as the regulated "
              "Castle Rock peak of the same probability plus a local "
              "contribution:")
    para(doc, "Q(p) = Q_reg(p) + Q_unreg(p) × (DA − 2,229)/2,229 × 0.80",
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "The local term scales off the UNREGULATED curve, not the "
              "regulated one. All three tributaries enter below Mossyrock "
              "and Riffe and respond to the storm rather than to the "
              "release, so the storm-scale curve is what governs them; "
              "scaling them off the regulated curve would shrink them in "
              "proportion to a reservoir upstream of them, about 18 percent "
              "low at the 1,000-year event.")
    para(doc, "Pairing the unregulated and regulated curves at a common "
              "probability is not a coincident-frequency assumption. It is "
              "the same river in the same event, and the regulated curve was "
              "derived from the unregulated one by routing (Section 5.4). "
              "Nothing in this method requires a tributary to be "
              "simultaneously at its own 1,000-year event, which is the "
              "assumption the earlier same-probability summation methods "
              "carried and could not defend.")

    h2(doc, "8.2 Local Contribution by Drainage Area")
    para(doc, "Drainage areas here are taken from the StreamStats "
              "delineations prepared for this analysis, which give 2,229 "
              "square miles at the Castle Rock gage. Section 2 cites 2,238 "
              "square miles from the earlier memoranda. The two differ by "
              "0.4 percent and the difference is immaterial to a ratio, but "
              "the delineated value is used throughout this section so that "
              "every area in it comes from one consistent source.")
    para(doc, "Each site uses the full incremental area between it and the "
              "gage, not the named tributary basin alone. The difference is "
              "ungaged local drainage along both banks, which contributes "
              "whether or not it carries a name — 247 square miles by "
              "incremental area against 197.5 for the three named basins.")
    table(doc,
          ["Location", "Drainage area (sq mi)", "Incremental (sq mi)",
           "Local as % of gage"],
          [["Castle Rock gage", "2,229", "—", "—"],
           ["Below Arkansas Creek", "2,278", "49", "2.2%"],
           ["Below Ostrander Creek", "2,335", "106", "4.8%"],
           ["Below Coweeman River", "2,476", "247", "11.1%"]],
          align_right=(1, 2, 3))
    para(doc, "A plain drainage-area ratio is adopted, with no unit-runoff "
              "adjustment. Three independent lines of evidence support that "
              "at the magnitudes that matter. Paired against the routed "
              "unregulated Castle Rock peak for the 23 water years in which "
              "both rivers peaked in the same storm, the Coweeman gage "
              "record (USGS 14245000, WY1950–1996) runs about 1.5 times its "
              "area share at common events but converges monotonically "
              "toward parity as events grow; in February 1996 — a 212,245 "
              "cfs unregulated event, 92 percent of the 1,000-year flow — it "
              "sat at 1.04 times its area share. The shorter Ecology record "
              "at gage 26C075 (WY2007–2019) gives 1.11 times for its largest "
              "bin, and that figure is a lower bound because the gage's "
              "rating ceiling censors precisely its largest events.")
    para(doc, "Basin mean precipitation tests the equal-depth assumption "
              "the area ratio rests on directly. Over the %d years %d–%d, "
              "PRISM annual precipitation averaged across the delineated "
              "boundaries gives a Coweeman-to-Castle-Rock ratio of %.2f "
              "(median %.2f, interquartile range %.2f–%.2f). The two "
              "basins receive very nearly the same depth."
              % (len(prism), int(prism["year"].min()), int(prism["year"].max()),
                 prism["ratio"].mean(), prism["ratio"].median(),
                 prism["ratio"].quantile(0.25), prism["ratio"].quantile(0.75)))
    para(doc, "Read together these lines resolve into one consistent "
              "picture. Precipitation is essentially equal per unit area, so "
              "the Coweeman's excess at common events is a RESPONSE "
              "difference — a small, low, steep basin concentrating runoff "
              "faster than a 2,229 square mile mainstem — rather than a "
              "precipitation difference. Response differences shrink as both "
              "basins saturate, which is why the flow ratio falls toward "
              "parity as events grow, and why at the largest event on record "
              "it converges on very nearly the precipitation ratio itself. "
              "A plain area ratio is therefore the right form at the "
              "magnitudes this study is concerned with, even though it would "
              "understate the tributary at ordinary flows.")

    h2(doc, "8.3 Timing Adjustment")
    para(doc, "The tributaries do not crest when the regulated Cowlitz "
              "does. They peak first and are already receding when the "
              "regulated mainstem arrives, so only part of each tributary "
              "peak is present at the moment that governs the combined "
              "flow. Measuring the Coweeman flow at the hour of the "
              "regulated Castle Rock crest, as a fraction of its own peak "
              "in the same storm, over 78 events:")
    table(doc,
          ["Castle Rock unregulated peak", "Events", "Median ratio"],
          [["20,000–40,000 cfs", "52", "0.806"],
           ["40,000–60,000 cfs", "19", "0.781"],
           ["Above 60,000 cfs", "7", "0.413"],
           ["All events", "78", "0.789"]],
          align_right=(1, 2))
    para(doc, "A factor of 0.80 is adopted. It is where the two "
              "well-sampled bins group, and it is the value used in the "
              "2009 restudy, so it carries precedent as well as data. It is "
              "very likely conservative: a higher timing factor adds more "
              "local flow, and the seven largest events in the record sit "
              "well below it.")
    para(doc, "The lower figure from those seven largest events is not "
              "adopted, for three reasons. The sample is small. Its "
              "mechanism is tributary lead time rather than event magnitude "
              "— the ratio correlates strongly with how far ahead the "
              "Coweeman crests (Spearman rho −0.55, p < 0.0001) while lead "
              "time itself shows no significant relationship to event size "
              "(rho −0.04, p = 0.73), so the low tail figure is a property "
              "of those particular storms rather than of large events "
              "generally. And three events whose Coweeman crest exceeded the "
              "gage rating are absent from that bin altogether — events in "
              "which the tributary was at its largest, so their absence "
              "biases the figure downward.")

    h2(doc, "8.4 Results")
    para(doc, "Regulated peak flows at the four locations are given below "
              "for the probabilities of primary interest; complete ordinates "
              "with confidence limits are in Appendix F.")
    # Built from the CSV rather than typed in: a hand-keyed summary table
    # silently goes stale the moment the source script is re-run with a
    # different lag factor or drainage area.
    headline = [0.10, 0.02, 0.01, 0.005, 0.002, 0.001]
    r84 = []
    for aep in headline:
        row = conf.iloc[(conf["AEP"] - aep).abs().idxmin()]
        r84.append(["%g%%" % (row["AEP"] * 100),
                    fmt(row["castle_rock_gage_cfs"]),
                    fmt(row["below_arkansas_creek_cfs"]),
                    fmt(row["below_ostrander_creek_cfs"]),
                    fmt(row["below_coweeman_river_cfs"])])
    table(doc,
          ["AEP", "Castle Rock gage", "Below Arkansas", "Below Ostrander",
           "Below Coweeman"],
          r84, align_right=(1, 2, 3, 4))
    tgt = conf.iloc[(conf["AEP"] - 0.001).abs().idxmin()]
    para(doc, "At the 1,000-year event the local contribution below the "
              "Coweeman confluence is %s cfs, raising the regulated peak "
              "%.1f percent above the gage. The increase is proportionally "
              % (fmt(tgt["below_coweeman_river_local_cfs"]),
                 100 * (tgt["below_coweeman_river_cfs"]
                        - tgt["castle_rock_gage_cfs"])
                 / tgt["castle_rock_gage_cfs"]) +
              "largest in the middle of the range, near the 2 percent "
              "event, and falls away in the extreme tail: the regulated "
              "curve steepens there as the projects lose the ability to "
              "hold back a rare inflow, while the local term grows only with "
              "the flatter unregulated curve.")
    figure(doc, "below_conf",
           "Figure 8-1. Regulated peak flow frequency at all four "
           "locations. The local contribution at each site is the "
           "incremental drainage area scaled off the unregulated curve and "
           "reduced by the 0.80 timing factor. The lower panel gives each "
           "site as a percentage increase over the gage.")
    figure(doc, "freq_gage",
           "Figure 8-2. Castle Rock gage, 2,229 square miles: the "
           "unregulated and regulated curves of Sections 4 and 5, with the "
           "regulated 95 percent confidence band from Section 5.6. This is "
           "the pair every downstream location is built from — the "
           "regulated curve sets the base and the unregulated curve drives "
           "the local contribution.")
    figure(doc, "freq_arkansas",
           "Figure 8-3. Below Arkansas Creek, 2,278 square miles.")
    figure(doc, "freq_ostrander",
           "Figure 8-4. Below Ostrander Creek, 2,335 square miles.")
    figure(doc, "freq_coweeman",
           "Figure 8-5. Below the Coweeman River, 2,476 square miles.")

    h2(doc, "8.5 Behavior of the Transform at Extreme Flows")
    para(doc, "A reservoir cannot remove volume from a flood, only move it "
              "in time — at the 5-day duration the regulated to unregulated "
              "ratio already exceeds one (Section 5.2). Riffe Lake holds "
              "358,116 acre-feet between the winter rule curve at 745.5 "
              "feet and full pool at 778.5 feet, which is 180,550 cfs-days. "
              "Once an event's pre-crest inflow fills that, every further "
              "cubic foot per second passes through and the regulated peak "
              "equals the unregulated one.")
    para(doc, "The adopted transform already shows this turning. Its "
              "maximum reduction is 67,558 cfs near the 0.2 percent event "
              "and it declines beyond that, because larger events fill the "
              "pool sooner. That maximum is a check on itself: 180,550 "
              "cfs-days divided by 67,558 cfs is 2.7 days, a plausible time "
              "from flood onset to crest in this basin, so the turn in the "
              "transform and the reservoir's storage volume corroborate one "
              "another.")
    figure(doc, "convergence",
           "Figure 8-6. Regulated against unregulated peak at Castle Rock, "
           "with the transform drawn converging on the 1:1 line. The "
           "convergence limb is an ESTIMATE drawn for illustration. It is "
           "not fitted, is used in no result in this memorandum, and does "
           "not affect any value in Sections 8.1 through 8.4 or in the "
           "appendices.")
    para(doc, "The convergence point is drawn rather than fitted because "
              "fitting does not survive the data. The 48 synthetic members "
              "cluster between 265,000 and 280,000 cfs unregulated, and "
              "their reductions there span 1,869 to 99,478 cfs — the "
              "December 1933 member passes 99 percent of its inflow while "
              "the December 2015 member, at almost the same magnitude, is "
              "still holding back 98,000 cfs. Convergence depends on "
              "hydrograph shape and starting pool, not on peak magnitude "
              "alone, so a regression through those points is not "
              "meaningful: fitting all 48 places the crossing at 931,600 "
              "cfs and fitting the largest twelve places it at 259,600. The "
              "value shown, 500,000 cfs, is consistent with the adopted "
              "transform's own declining reduction, which extrapolates to "
              "zero near 480,000 to 500,000 cfs. It lies beyond the "
              "10,000-year unregulated event of 374,643 cfs and therefore "
              "outside the range of every result presented here. Settling "
              "it as a number rather than a drawing would require scaling "
              "additional events through the reservoir model.")
    para(doc, "The confidence band shown at each downstream location is "
              "the Castle Rock regulated band of Section 5.6, carried "
              "forward and translated by that site's local contribution. No "
              "additional uncertainty is added for the local term itself. "
              "That term carries real uncertainty, but it is a few percent "
              "of a quantity whose own band is already wider than the flow "
              "it brackets, so adding it would not be visible on these "
              "figures and would imply a precision this method does not "
              "have. The band below the "
              "Coweeman confluence spans %s to %s cfs at the "
              "1,000-year event. The choice of timing factor is small "
              % (fmt(tgt["below_coweeman_river_lower_cfs"]),
                 fmt(tgt["below_coweeman_river_upper_cfs"])) +
              "against that: the full 0.41 to 1.00 range moves the total "
              "7.2 percent, about 6 percent of the width of the confidence "
              "band.")

    h1(doc, "Appendix A. Assembled Unregulated Record by Water Year")
    para(doc, "Flows are in cubic feet per second. Source abbreviations: "
              "USGS peak, the observed instantaneous annual peak; USGS "
              "daily, rolling maxima of the observed daily record; "
              "Holdout, the hourly unregulated record from the Mossyrock "
              "holdout; Regression, the two-day storage-change relation; "
              "Mass balance, the daily unregulated series.")
    a_rows = [tuple(c if c else "—" for c in r) for r in APPENDIX_A_ROWS]
    table(doc, ["WY", "Peak", "Peak source", "1-Day", "1-Day source",
               "3-Day", "5-Day", "Duration source"], a_rows,
        widths=[700, 1200, 1360, 1200, 1360, 1100, 1100, 1340],
        align_right=(0, 1, 3, 5, 6), font_size=16)
    caption(doc, "Table A-1. Assembled unregulated record, water years "
                "1927 through 2026.")

    h1(doc, "Appendix B. Unregulated Frequency Curve Ordinates")
    para(doc, "Computed curve, expected probability curve, variance of the "
              "log, and 5 and 95 percent confidence limits for every "
              "ordinate reported by HEC-SSP. Flows are in cubic feet per "
              "second.")
    table(doc, ["Duration", "AEP", "Variance of log", "z", "Computed",
               "5% limit", "95% limit", "Expected probability"],
        APPENDIX_B_ROWS,
        widths=[1200, 1000, 1300, 700, 1400, 1400, 1400, 960],
        align_right=(1, 2, 3, 4, 5, 6, 7), font_size=16)
    caption(doc, "Table B-1. Frequency curve ordinates, all four "
                "durations.")

    h1(doc, "Appendix C. Unregulated Frequency Curves by Duration")
    para(doc, "Individual curves as produced by HEC-SSP: computed curve, "
              "expected probability curve, 5 and 95 percent confidence "
              "limits, and observed events at Hirsch-Stedinger plotting "
              "positions. The peak figure reflects the censored treatment "
              "described in Section 4.8.")
    figure(doc, "curve_peak", "Figure C-1. Unregulated instantaneous peak "
                             "frequency curve.")
    figure(doc, "curve_1day", "Figure C-2. Unregulated 1-day frequency "
                             "curve.")
    figure(doc, "curve_3day", "Figure C-3. Unregulated 3-day frequency "
                             "curve.")
    figure(doc, "curve_5day", "Figure C-4. Unregulated 5-day frequency "
                             "curve.")

    h1(doc, "Appendix D. Adjusted Regulated Peak Record by Water Year")
    para(doc, "Observed, simulated and adjusted regulated peaks for every "
              "shared water year, with the screening result recorded for "
              "each. Flows are in cubic feet per second.")
    table(doc, ["WY", "Observed", "WCM_RC", "Obs_RC", "Difference",
               "Adjusted", "Screen"], APPENDIX_D_ROWS,
        widths=[900, 1300, 1300, 1300, 1300, 1300, 1960],
        align_right=(0, 1, 2, 3, 4, 5), font_size=16)
    caption(doc, "Table D-1. Adjusted regulated peak record, water years "
                "1974 through 2024.")

    h1(doc, "Appendix E. Regulated Frequency Curve Ordinates and "
            "Uncertainty Terms")
    para(doc, "The adopted regulated curve with its uncertainty band "
              "(Section 5.6, FREQ_TERM_MODE = %r), and the terms that "
              "produced it, for every ordinate. Flows are in cubic feet "
              "per second." % freq_mode)
    e_rows = []
    for _, r in reg.sort_values("AEP", ascending=False).iterrows():
        e_rows.append([
            "%.4f" % r["AEP"], fmt(r["unreg_expected_cfs"]),
            fmt(r["reg_inferred_cfs"]), fmt(r["reg_lower_95pct_cfs"]),
            fmt(r["reg_upper_95pct_cfs"]), "%.2f" % r["transform_slope_b"],
            "%.1f" % r["reduction_pct"],
            "yes" if r["extrapolated"] else "no",
        ])
    table(doc, ["AEP", "Unregulated", "Regulated", "Lower 95%", "Upper 95%",
               "Local slope", "Reduction %", "Extrap."], e_rows,
        widths=[900, 1300, 1300, 1300, 1300, 1160, 1200, 900],
        align_right=(0, 1, 2, 3, 4, 5, 6))
    caption(doc, "Table E-1. Regulated frequency curve ordinates and "
                "uncertainty terms.")

    h1(doc, "Appendix F. Regulated Frequency Ordinates by Location")
    para(doc, "Regulated peak flow at each of the four locations of Section "
              "8, with 95 percent confidence limits. Flows are in cubic "
              "feet per second. At the gage the unregulated curve is shown "
              "alongside; downstream, the local contribution is shown, "
              "being the incremental drainage area scaled off the "
              "unregulated curve and reduced by the 0.80 timing factor. "
              "Confidence limits at the three downstream locations are the "
              "gage's own, translated by the local contribution — see "
              "Section 8.4.")

    site_tables = [
        ("Castle Rock gage", 2229, "castle_rock_gage", "Unregulated",
         "cowlitz_unreg_cfs"),
        ("Below Arkansas Creek", 2278, "below_arkansas_creek", "Local",
         "below_arkansas_creek_local_cfs"),
        ("Below Ostrander Creek", 2335, "below_ostrander_creek", "Local",
         "below_ostrander_creek_local_cfs"),
        ("Below Coweeman River", 2476, "below_coweeman_river", "Local",
         "below_coweeman_river_local_cfs"),
    ]
    for n, (label, area, key, second, second_col) in enumerate(site_tables, start=1):
        rows = []
        for _, r in conf.iterrows():
            rows.append(["%.4f" % r["AEP"],
                         fmt(r[second_col]),
                         fmt(r["%s_cfs" % key]),
                         fmt(r["%s_lower_cfs" % key]),
                         fmt(r["%s_upper_cfs" % key])])
        table(doc, ["AEP", second, "Regulated", "Lower 95%", "Upper 95%"],
              rows, widths=[1400, 1800, 1800, 1800, 1800],
              align_right=(0, 1, 2, 3, 4))
        caption(doc, "Table F-%d. %s, %s square miles." % (n, label, "{:,}".format(area)))

    # ============================================================
    # REFERENCES
    # ============================================================
    h1(doc, "References")
    for ref in [
        "England, J.F., Jr., Cohn, T.A., Faber, B.A., Stedinger, J.R., "
        "Thomas, W.O., Jr., Veilleux, A.G., Kiang, J.E., and Mason, R.R., "
        "Jr., 2019, Guidelines for determining flood flow frequency - "
        "Bulletin 17C (ver. 1.1): U.S. Geological Survey Techniques and "
        "Methods, book 4, chap. B5, 148 p.",
        "Lind, G.D., Lamontagne, J.R., and Stonewall, A.J., 2020, "
        "Development of regional skew coefficients for selected flood "
        "durations in the Columbia River Basin, northwestern United "
        "States and British Columbia, Canada (ver. 1.1, October 2020): "
        "U.S. Geological Survey Scientific Investigations Report "
        "2020-5073, 48 p.",
        "Mastin, M.C., Konrad, C.P., Veilleux, A.G., and Tecca, A.E., "
        "2016, Magnitude, frequency, and trends of floods at gaged and "
        "ungaged sites in Washington, based on data through water year "
        "2014 (ver. 1.2, November 2017): U.S. Geological Survey "
        "Scientific Investigations Report 2016-5118, 70 p.",
        "U.S. Army Corps of Engineers, 2025, Risk Assessment for Flood "
        "Risk Management Studies, EM 1110-2-1619, 29 September 2025.",
        "U.S. Army Corps of Engineers, Portland District, 2009, Cowlitz "
        "River Hydrology Restudy.",
        "U.S. Army Corps of Engineers, 2014, Water Control Manual, "
        "Mossyrock Project, Cowlitz River, Washington.",
        "U.S. Army Corps of Engineers, Hydrologic Engineering Center, "
        "HEC-ResSim Reservoir System Simulation, and HEC-SSP Statistical "
        "Software Package.",
        "U.S. Geological Survey, National Water Information System: "
        "streamflow and peak-flow records for station 14243000, Cowlitz "
        "River at Castle Rock, Washington.",
    ]:
        para(doc, ref, after=100)

    REVIEW_NOTES.append(
        "Two source memos cite the Mossyrock/Mayfield Water Control Manual "
        "with different authors -- USACE (unregulated memo) vs. Tacoma "
        "Power (regulated memo). The combined reference list keeps the "
        "USACE attribution from the unregulated memo; confirm which is "
        "correct before this goes out.")

    doc.save(OUT_DOCX)
    print("Wrote %s" % OUT_DOCX)
    if MISSING:
        print("\nMISSING FIGURES:")
        for m in MISSING:
            print("  -", m)
    if REVIEW_NOTES:
        print("\nREVIEW NOTES:")
        for n in REVIEW_NOTES:
            print("  -", n)


if __name__ == "__main__":
    main()
